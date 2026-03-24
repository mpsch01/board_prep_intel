"""
csv_to_docx.py
--------------
Converts ABFM_ITE_Enriched.csv into a compact, formatted DOCX question bank.

Usage:
    python csv_to_docx.py

Output:
    04_outputs/ABFM_ITE_Question_Bank_v6.docx

v6 changes vs v5:
  - TWO-COLUMN newspaper page layout  (~halves page count on its own)
  - Per-question compact table with 2-column answer choices (A|B, C|D, E)
  - Table auto-sizes to column width (no hardcoded Inches)
  - Thin outer table border; barely-visible horizontal inside rows; no vertical
  - Minimal cell padding throughout
  - No separate title page (compact header only)
"""

import re
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# PATHS
# ---------------------------------------------------------------------------
BASE        = Path(r"C:\Users\mpsch\Desktop\claude_knowledge\board_prep")
CSV_PATH    = BASE / "ite_exam/03_database/ABFM_ITE_Enriched.csv"
OUTPUT_DOCX = BASE / "ite_exam/04_outputs/ABFM_ITE_Question_Bank_v6.docx"

# ---------------------------------------------------------------------------
# COLORS
# ---------------------------------------------------------------------------
GRAY_HEADER  = "D9D9D9"
LABEL_COLOR  = RGBColor(0x33, 0x33, 0x33)
ANSWER_COLOR = RGBColor(0x1A, 0x56, 0x76)   # dark teal for "Correct Answer:"

# ---------------------------------------------------------------------------
# DOCUMENT-LEVEL SETUP
# ---------------------------------------------------------------------------
def configure_two_column(doc):
    """Apply 2-column newspaper layout to every section."""
    for section in doc.sections:
        sectPr = section._sectPr
        existing = sectPr.find(qn("w:cols"))
        if existing is not None:
            sectPr.remove(existing)
        cols = OxmlElement("w:cols")
        cols.set(qn("w:num"),        "2")
        cols.set(qn("w:space"),      "360")   # 0.25 in gap (twips)
        cols.set(qn("w:equalWidth"), "1")
        sectPr.append(cols)

# ---------------------------------------------------------------------------
# TABLE HELPERS
# ---------------------------------------------------------------------------
def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=1, bottom=1, left=3, right=3):
    """All values in points; converted to twips (×20) internally."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"),    str(int(val * 20)))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _get_or_add_tblPr(tbl):
    """Return tblPr element, creating it as first child if absent."""
    tbl_pr = tbl._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tbl_pr)
    return tbl_pr


def set_table_full_width(tbl):
    """Force table to 100 % of the text column (works in multi-column docs)."""
    tbl_pr = _get_or_add_tblPr(tbl)
    tbl_w  = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"),    "5000")   # 5000 × 0.02 % = 100 %


def set_table_borders(tbl):
    """Thin outer border; faint horizontal inside rows; no vertical divider."""
    tbl_pr  = _get_or_add_tblPr(tbl)
    tbl_bdr = OxmlElement("w:tblBorders")

    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")        # 0.5 pt
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "AAAAAA")
        tbl_bdr.append(b)

    ins_h = OxmlElement("w:insideH")
    ins_h.set(qn("w:val"),   "single")
    ins_h.set(qn("w:sz"),    "2")        # 0.25 pt — barely visible
    ins_h.set(qn("w:space"), "0")
    ins_h.set(qn("w:color"), "CCCCCC")
    tbl_bdr.append(ins_h)

    ins_v = OxmlElement("w:insideV")
    ins_v.set(qn("w:val"),   "none")
    ins_v.set(qn("w:sz"),    "0")
    ins_v.set(qn("w:space"), "0")
    ins_v.set(qn("w:color"), "auto")
    tbl_bdr.append(ins_v)

    tbl_pr.append(tbl_bdr)


def _p_tight(cell, line_spacing_pt=None):
    """Return the paragraph of a cell with zero before/after spacing."""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if line_spacing_pt:
        p.paragraph_format.line_spacing = Pt(line_spacing_pt)
    return p


def add_full_row(tbl, text, size=9.5, bold=False, italic=False,
                 color=None, bg=None,
                 label=None, label_color=None,
                 top=1, bottom=1, line_spacing=None):
    """Append a full-width (merged) row to tbl."""
    row  = tbl.add_row()
    cell = row.cells[0].merge(row.cells[1])
    set_cell_margins(cell, top=top, bottom=bottom, left=3, right=3)
    if bg:
        set_cell_shading(cell, bg)
    p = _p_tight(cell, line_spacing)
    if label:
        rl = p.add_run(label)
        rl.font.size = Pt(8)
        rl.font.bold = True
        if label_color:
            rl.font.color.rgb = label_color
    if text:
        r = p.add_run(text)
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = color


def add_choice_row(tbl, left_letter, left_text,
                   right_letter=None, right_text=None):
    """Append a 2-column answer-choice row (A|B or C|D or E|empty)."""
    row = tbl.add_row()
    for i, (letter, text) in enumerate(
            [(left_letter, left_text), (right_letter, right_text)]):
        cell = row.cells[i]
        set_cell_margins(cell, top=0, bottom=0, left=4, right=3)
        p = _p_tight(cell, line_spacing_pt=11)
        if letter and text:
            r = p.add_run(f"{letter})  {text}")
            r.font.size = Pt(9.5)

# ---------------------------------------------------------------------------
# MISC HELPERS
# ---------------------------------------------------------------------------
def add_page_break(doc):
    p   = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def set_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

# ---------------------------------------------------------------------------
# STEM PARSER  (unchanged)
# ---------------------------------------------------------------------------
def parse_stem(raw_stem):
    """Split QuestionStem blob into (scenario_text, {letter: choice_text})."""
    if not isinstance(raw_stem, str):
        return "", {}
    m = re.search(r"\n(?=A\))", raw_stem)
    if not m:
        return raw_stem.strip(), {}
    scenario      = raw_stem[:m.start()].strip()
    choices_block = raw_stem[m.start():].strip()
    choices = {}
    current_letter = None
    current_lines  = []
    for line in choices_block.split("\n"):
        mc = re.match(r"^([A-E])\)\s*(.*)", line)
        if mc:
            if current_letter is not None:
                text = "\n".join(current_lines).strip()
                if text:
                    choices[current_letter] = text
            current_letter = mc.group(1)
            current_lines  = [mc.group(2)] if mc.group(2).strip() else []
        elif current_letter is not None and line.strip():
            current_lines.append(line)
    if current_letter is not None:
        text = "\n".join(current_lines).strip()
        if text:
            choices[current_letter] = text
    return scenario, choices

# ---------------------------------------------------------------------------
# WRITE ONE QUESTION BLOCK
# ---------------------------------------------------------------------------
def write_question(doc, row, global_num):
    qid         = str(row.get("Question ID",   ""))
    year        = str(row.get("ExamYear",       ""))
    category    = str(row.get("PrimaryCategory",""))
    correct     = str(row.get("CorrectAnswer",  "")).strip().upper()
    raw_stem    = row.get("QuestionStem", "")
    explanation = str(row.get("Explanation",    "")).strip()

    ref = str(row.get("References", "")).strip()
    if not ref or ref.lower() in ("nan", ""):
        ref = "No reference provided."

    scenario, choices = parse_stem(raw_stem)

    # ── One table per question ──────────────────────────────────────────────
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    set_table_full_width(tbl)
    set_table_borders(tbl)

    # Row 1 — gray header bar
    add_full_row(tbl,
                 text=f"#{global_num}   |   {qid}   |   {year}   |   {category}",
                 size=8, bold=True, bg=GRAY_HEADER,
                 top=1, bottom=1)

    # Row 2 — question stem
    add_full_row(tbl, text=scenario,
                 size=9.5, top=3, bottom=2, line_spacing=12)

    # Rows 3+ — answer choices, 2 per row (A|B, C|D, E|—)
    letters = [l for l in ("A", "B", "C", "D", "E") if l in choices]
    for i in range(0, len(letters), 2):
        left_l  = letters[i]
        right_l = letters[i + 1] if i + 1 < len(letters) else None
        add_choice_row(tbl,
                       left_l,  choices[left_l],
                       right_l, choices.get(right_l) if right_l else None)

    # Correct answer row
    add_full_row(tbl,
                 text=f"Correct Answer: {correct}",
                 size=9.5, bold=True, color=ANSWER_COLOR,
                 top=2, bottom=1)

    # Explanation row
    if explanation and explanation.lower() not in ("nan", ""):
        add_full_row(tbl,
                     text=explanation, size=9.5,
                     label="EXPLANATION  ", label_color=LABEL_COLOR,
                     top=2, bottom=1, line_spacing=12)

    # Reference row
    add_full_row(tbl,
                 text=ref, size=9, italic=True,
                 label="REFERENCE  ", label_color=LABEL_COLOR,
                 top=2, bottom=2)

    # Spacer paragraph between questions (very small)
    p_gap = doc.add_paragraph()
    p_gap.paragraph_format.space_before = Pt(0)
    p_gap.paragraph_format.space_after  = Pt(3)

# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------
def main():
    print("Loading data...")
    df = pd.read_csv(CSV_PATH, dtype=str)
    df["_year"] = df["ExamYear"].str.strip()
    df["_qid"]  = df["Question ID"].str.strip()
    df = df.sort_values(by=["_year", "_qid"]).reset_index(drop=True)
    years = sorted(df["_year"].unique())

    print("Building DOCX...")
    doc = Document()

    # ── Tight margins ───────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.3)
        section.bottom_margin = Cm(1.3)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)

    # ── 2-column layout ─────────────────────────────────────────────────────
    configure_two_column(doc)

    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9.5)

    # ── Compact document header (no separate title page) ────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("ABFM ITE Question Bank")
    set_font(r, 13, bold=True)
    r.font.name = "Aptos"

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run(
        f"ITE  {years[0]}\u2013{years[-1]}  \u2022  {len(df):,} Questions"
    )
    set_font(r2, 9)
    r2.font.name = "Aptos"

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Questions by year ───────────────────────────────────────────────────
    global_num = 1
    for year in years:
        year_df = df[df["_year"] == year]

        heading = doc.add_heading(f"{year}  ({len(year_df)} questions)", level=1)
        heading.paragraph_format.space_before = Pt(4)
        heading.paragraph_format.space_after  = Pt(2)
        for run in heading.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(11)

        for _, row in year_df.iterrows():
            write_question(doc, row, global_num)
            global_num += 1

        if year != years[-1]:
            add_page_break(doc)

    # ── Save ────────────────────────────────────────────────────────────────
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    total = global_num - 1
    print(f"\nDone: {total} questions written.")
    print(f"Output: {OUTPUT_DOCX}")
    print("\nYear breakdown:")
    for year in years:
        n = len(df[df["_year"] == year])
        print(f"  {year}:  {n} questions")


if __name__ == "__main__":
    main()
