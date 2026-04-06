"""
word_doc_defaults.py — Default Word Document Style for Mikey's Project
=======================================================================
St. Luke's / ITE Intelligence color scheme + Aptos font.

Usage in any python-docx build script:
    from word_doc_defaults import *   # or import specific helpers

CLAUDE: When building any Word document for this project, import and
apply these defaults. See CLAUDE.md → "Default Word Doc Style" section.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ==============================================================================
# COLOR PALETTE  (St. Luke's — dark navy / gold / medium blue)
# ==============================================================================

NAVY       = "1B3564"    # Primary: titles, headers, section labels
GOLD       = "C8922A"    # Accent: left border bars, highlights
BLUE       = "2E5F9C"    # Secondary: subheaders, labels
LIGHT_BLUE = "EBF0F7"    # Background: answer boxes, shaded cells
DARK_TEXT  = "1A1A2E"    # Body text (near-black)
MED_GRAY   = "808080"    # Secondary text, footers, captions
GREEN      = (62, 141, 39)   # Correct-answer green — use as RGBColor(*GREEN)

# Convenience RGBColor objects
RGB_NAVY      = RGBColor(0x1B, 0x35, 0x64)
RGB_GOLD      = RGBColor(0xC8, 0x92, 0x2A)
RGB_BLUE      = RGBColor(0x2E, 0x5F, 0x9C)
RGB_DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
RGB_GRAY      = RGBColor(0x80, 0x80, 0x80)
RGB_GREEN     = RGBColor(62, 141, 39)

# ==============================================================================
# TYPOGRAPHY
# ==============================================================================

DEFAULT_FONT   = "Aptos"
FONT_TITLE     = 26       # pt — document title
FONT_SUBTITLE  = 14       # pt — subtitle line
FONT_HEADING   = 12       # pt — section/question headers
FONT_BODY      = 11       # pt — body text, answer choices
FONT_SMALL     = 10       # pt — explanations, captions, footer
FONT_TINY      = 9        # pt — fine print

# ==============================================================================
# PAGE SETUP  (US Letter, 1" margins)
# ==============================================================================

def apply_page_setup(doc):
    """Apply US Letter size and 1-inch margins to all sections."""
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

def set_default_font(doc):
    """Set Aptos as the document-wide default font."""
    style = doc.styles['Normal']
    style.font.name = DEFAULT_FONT
    style.font.size = Pt(FONT_BODY)

def new_document():
    """Return a new Document with page setup and default font applied."""
    doc = Document()
    apply_page_setup(doc)
    set_default_font(doc)
    return doc

# ==============================================================================
# PARAGRAPH HELPERS
# ==============================================================================

def add_shading(paragraph, fill_color):
    """Apply background fill color to a paragraph (use hex string, e.g. 'EBF0F7')."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_color)
    pPr.append(shd)

def add_left_border(paragraph, color=GOLD, size=24):
    """Add a thick vertical left border bar (gold accent by default)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    str(size))
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)

def add_bottom_border(paragraph, color="CCCCCC", size=4):
    """Add a thin horizontal rule below a paragraph (divider line)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_paragraph_format(para, keep_with_next=False, widow_control=True, keep_together=False):
    """Apply keep_with_next, widow control, and keep_lines_together."""
    pPr = para._p.get_or_add_pPr()
    if widow_control:
        wc = OxmlElement('w:widowControl')
        wc.set(qn('w:val'), '1')
        pPr.append(wc)
    if keep_with_next:
        pPr.append(OxmlElement('w:keepNext'))
    if keep_together:
        pPr.append(OxmlElement('w:keepLines'))

# ==============================================================================
# STANDARD PARAGRAPH STYLES
# ==============================================================================

def add_title(doc, text):
    """Document title — 26pt bold navy, centered, with thin navy rule below."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name  = DEFAULT_FONT
    run.font.size  = Pt(FONT_TITLE)
    run.font.bold  = True
    run.font.color.rgb = RGB_NAVY
    add_bottom_border(p, color=NAVY, size=12)
    set_paragraph_format(p, keep_with_next=True)
    return p

def add_subtitle(doc, text):
    """Subtitle line — 14pt blue, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name  = DEFAULT_FONT
    run.font.size  = Pt(FONT_SUBTITLE)
    run.font.color.rgb = RGB_BLUE
    set_paragraph_format(p, keep_with_next=True)
    return p

def add_section_header(doc, text, level=1):
    """
    Section header with gold left border + light blue background.
    level 1 = 12pt bold navy  (question headers, major sections)
    level 2 = 11pt bold blue  (subheadings)
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    add_left_border(p, color=GOLD, size=24)
    add_shading(p, "EFF3FA")
    set_paragraph_format(p, keep_with_next=True)
    run = p.add_run(text)
    run.font.name = DEFAULT_FONT
    if level == 1:
        run.font.size = Pt(FONT_HEADING)
        run.font.bold = True
        run.font.color.rgb = RGB_NAVY
    else:
        run.font.size = Pt(FONT_BODY)
        run.font.bold = True
        run.font.color.rgb = RGB_BLUE
    return p

def add_body_text(doc, text, indent=0.25):
    """Standard body paragraph — 11pt dark text, optional left indent."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    set_paragraph_format(p, widow_control=True)
    run = p.add_run(text)
    run.font.name  = DEFAULT_FONT
    run.font.size  = Pt(FONT_BODY)
    run.font.color.rgb = RGB_DARK_TEXT
    return p

def add_answer_choice(doc, letter, text, indent=0.25):
    """
    Single answer choice — 11pt dark text, regular color (no green).
    letter: e.g. 'A', 'B', 'C', 'D', 'E'
    text:   the answer option text
    Format: "A.  Some answer text here"
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_after  = Pt(2)
    set_paragraph_format(p, widow_control=True)
    label_run = p.add_run(f"{letter}.  ")
    label_run.font.name      = DEFAULT_FONT
    label_run.font.size      = Pt(FONT_BODY)
    label_run.font.bold      = True
    label_run.font.color.rgb = RGB_DARK_TEXT
    text_run = p.add_run(text)
    text_run.font.name      = DEFAULT_FONT
    text_run.font.size      = Pt(FONT_BODY)
    text_run.font.color.rgb = RGB_DARK_TEXT
    return p

def add_correct_answer_label(doc, letter, indent=0.25):
    """
    'CORRECT ANSWER: X' label — 11pt bold green (RGB 62, 141, 39).
    Use as the first line of the explanation block, before body text.
    letter: e.g. 'A', 'B', 'C', 'D', 'E'
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    set_paragraph_format(p, keep_with_next=True)
    run = p.add_run(f"CORRECT ANSWER: {letter}")
    run.font.name      = DEFAULT_FONT
    run.font.size      = Pt(FONT_BODY)
    run.font.bold      = True
    run.font.color.rgb = RGB_GREEN
    return p

def add_divider(doc):
    """Thin gray horizontal rule between sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    add_bottom_border(p, color="CCCCCC", size=4)
    return p

def add_page_number_footer(doc_section):
    """
    Add footer: left = doc title stub | right = page number (tab-aligned).
    Call once per section.
    """
    from docx.oxml.shared import OxmlElement as OE

    footer = doc_section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Tab stop at right margin for page number
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')   # right margin at 1" insets on 8.5" letter = 9360 DXA
    tabs.append(tab)
    pPr.append(tabs)

    left_run = p.add_run("2025 ABFM ITE  |  St. Luke's Family Medicine Residency")
    left_run.font.name = DEFAULT_FONT
    left_run.font.size = Pt(FONT_TINY)
    left_run.font.color.rgb = RGB_GRAY

    tab_run = p.add_run("\t")
    tab_run.font.size = Pt(FONT_TINY)

    page_run = p.add_run("Page ")
    page_run.font.name = DEFAULT_FONT
    page_run.font.size = Pt(FONT_TINY)
    page_run.font.color.rgb = RGB_GRAY

    # PAGE field
    fld_run = p.add_run()
    fld_run.font.size = Pt(FONT_TINY)
    fld_run.font.color.rgb = RGB_GRAY
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    fld_run._r.extend([fldChar1, instrText, fldChar2])

# ==============================================================================
# ENCODING REPAIR  (for CSV/text sources with Windows-1252 mojibake)
# ==============================================================================

def fix_encoding(text):
    """Fix UTF-8 bytes that were misread as cp1252 (Windows-1252)."""
    if not text:
        return text
    try:
        return text.encode('cp1252').decode('utf-8')
    except (UnicodeDecodeError, UnicodeEncodeError):
        return text

def clean_text(text):
    """Fix remaining mojibake patterns via direct string replacement."""
    if not text:
        return text
    replacements = [
        ('â€"', '–'), ('â€"', '—'), ('â€™', "'"), ('â€œ', '"'),
        ('â€', '"'), ('â‰¥', '≥'), ('â€°¥', '≥'), ('â€¢', '•'),
        ('Ã©', 'é'), ('â€˜', "'"), ('â€¦', '…'),
        ('\u00e2\u20ac\u201c', '–'), ('\u00e2\u20ac\u201d', '–'),
        ('\u00e2\u20ac\u2022', '•'), ('\u00e2\u20ac\u2018', '\u2018'),
        ('\u00e2\u20ac\u2019', '\u2019'), ('\u00e2\u20ac\u0153', 'œ'),
        ('\u00e2\u2030\xa5', '≥'),
    ]
    for bad, good in replacements:
        text = text.replace(bad, good)
    return text

def sanitize(text):
    """Convenience: apply both fix_encoding and clean_text."""
    return clean_text(fix_encoding(text))
