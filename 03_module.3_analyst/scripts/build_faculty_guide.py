#!/usr/bin/env python3
"""
build_faculty_guide.py — Generate ITE Report Guide for Faculty Advisors
========================================================================

Outputs: ../docs/ITE_Report_Guide_Faculty_v2.docx

Uses word_doc_defaults.py for St. Luke's color scheme, Aptos font, and helper functions.
Includes two custom callout styles: DATA SOURCE (green), LIMITATION (amber).

Usage:
    python3 build_faculty_guide.py
"""

from pathlib import Path
from word_doc_defaults import *
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ==============================================================================
# PATHS
# ==============================================================================

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent.parent
DOCS_DIR = SCRIPT_DIR.parent / "docs"
OUTPUT_PATH = DOCS_DIR / "ITE_Report_Guide_Faculty_v2.docx"

DOCS_DIR.mkdir(parents=True, exist_ok=True)

# ==============================================================================
# CUSTOM CALLOUT COLORS & HELPERS
# ==============================================================================

GREEN_HEX      = "276749"   # Data source border
AMBER_HEX      = "975A16"   # Limitation border
LIGHT_GREEN_BG = "F0F7F2"   # Light green background
LIGHT_AMBER_BG = "FEF9EC"   # Light amber background


def add_data_source_callout(doc, text):
    """📂 DATA SOURCE — green left border, light green background"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_left_border(p, color=GREEN_HEX, size=24)
    add_shading(p, LIGHT_GREEN_BG)
    label = p.add_run("📂 DATA SOURCE  ")
    label.font.name = DEFAULT_FONT
    label.font.size = Pt(FONT_BODY)
    label.font.bold = True
    label.font.color.rgb = RGBColor(0x27, 0x67, 0x49)
    body = p.add_run(text)
    body.font.name = DEFAULT_FONT
    body.font.size = Pt(FONT_BODY)
    body.font.color.rgb = RGB_DARK_TEXT
    return p


def add_limitation_callout(doc, text):
    """⚠ LIMITATION — amber left border, light amber background"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_left_border(p, color=AMBER_HEX, size=24)
    add_shading(p, LIGHT_AMBER_BG)
    label = p.add_run("⚠ LIMITATION  ")
    label.font.name = DEFAULT_FONT
    label.font.size = Pt(FONT_BODY)
    label.font.bold = True
    label.font.color.rgb = RGBColor(0x97, 0x5A, 0x16)
    body = p.add_run(text)
    body.font.name = DEFAULT_FONT
    body.font.size = Pt(FONT_BODY)
    body.font.color.rgb = RGB_DARK_TEXT
    return p


def add_bullet(doc, text, indent=0.5):
    """Bullet point with consistent formatting."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(FONT_BODY)
    run.font.color.rgb = RGB_DARK_TEXT
    return p


def add_data_table(doc, rows_data):
    """Two-column table: left col = label (bold navy), right col = description (dark text)."""
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    # Header row
    hdr = table.add_row()
    for i, h in enumerate(["Stage / Item", "What happens"]):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(FONT_SMALL)
        run.font.bold = True
        run.font.color.rgb = RGB_NAVY
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EFF3FA')
        tcPr.append(shd)

    # Data rows
    for label, desc in rows_data:
        row = table.add_row()
        lc = row.cells[0]
        lc.text = label
        lr = lc.paragraphs[0].runs[0]
        lr.font.name = DEFAULT_FONT
        lr.font.size = Pt(FONT_SMALL)
        lr.font.bold = True
        lr.font.color.rgb = RGB_NAVY
        rc = row.cells[1]
        rc.text = desc
        rr = rc.paragraphs[0].runs[0]
        rr.font.name = DEFAULT_FONT
        rr.font.size = Pt(FONT_SMALL)
        rr.font.color.rgb = RGB_DARK_TEXT

    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(5.4)

    doc.add_paragraph()  # spacer after table
    return table


# ==============================================================================
# BUILD DOCUMENT
# ==============================================================================

def build_faculty_guide():
    """Generate the faculty advisor methodology guide."""
    doc = new_document()

    # COVER PAGE
    add_title(doc, "ITE Score Analysis")
    add_subtitle(doc, "Faculty Advisor's Methodology Guide")
    add_subtitle(doc, "Data sources, derivation logic, and interpretation framework for program directors and faculty advisors")
    add_divider(doc)
    add_body_text(doc, "This guide accompanies the ITE Score Analysis report. It explains how each section of the report was generated, what data sources underlie it, and what the analytical limitations are. Faculty advisors using this report to guide resident advising should understand which findings are reliable and which require clinical judgment to interpret.")

    # =========================================================================
    # SECTION: SYSTEM OVERVIEW
    # =========================================================================
    add_section_header(doc, "System Overview", level=1)
    add_body_text(doc, "The report is generated by a multi-stage analysis process that draws from three primary sources: the resident's ABFM score report PDF, a structured database of ITE questions and clinical guidelines, and ABFM's published reference files. Each section of the report is produced independently from these inputs — understanding which source drives each section is the key to calibrating how much weight to place on its findings.")

    add_section_header(doc, "The central database", level=2)
    add_body_text(doc, "A central database serves as the single source of truth for all non-ABFM data in the report. It contains over 1,600 ITE questions (2018-2025), over 1,200 AAFP board review questions, approximately 2,000 clinical guideline articles, ICD-10 tags, vector embeddings, and article currency data. The pipeline runs per-resident — each score report PDF is processed independently. Longitudinal comparison requires a prior-year analysis from the same resident.")

    # =========================================================================
    # SECTION 0: EXAM AT A GLANCE
    # =========================================================================
    add_section_header(doc, "Section 0 — Exam at a Glance", level=1)
    add_data_source_callout(doc, "Exam specifications and national benchmarks are pulled from ABFM's published annual reference data. This includes: total items (200), scored item count, experimental item count, MPS (380), SEM (±38), administration count, and national mean scaled scores and standard deviations by PGY level. This file is updated annually after ABFM releases each year's score report. If the current year's data is not yet available, the most recent prior year is used as a fallback.")
    add_limitation_callout(doc, "National benchmarks are published at the program aggregate level. Individual PGY comparisons require official score report input — without it, the display falls back to the all-PGY national mean. If the current year's reference data is unavailable, displayed benchmarks reflect a prior year.")

    # =========================================================================
    # SECTION 1: SCORE DISPLAY
    # =========================================================================
    add_section_header(doc, "Section 1 — Score Display", level=1)
    add_section_header(doc, "How the three numbers are derived", level=2)
    add_body_text(doc, "Raw %: correct / total_scored — direct calculation from parsed exam data.")
    add_body_text(doc, "Scaled score has two sources. Official (labeled 'official'): extracted directly from the ABFM score report PDF. Estimated (labeled 'est.'): if no score report is provided, a linear approximation is used — scaled approximately equals 380 times (raw_pct / national_mean_pct). Accurate to ±10–15 points in the 330–480 range; less reliable at extremes.")
    add_body_text(doc, "Percentile: estimated from a normal distribution using the published national mean and SD for that PGY level. Not derived from ABFM's own percentile ranking.")

    add_section_header(doc, "Derived statistics", level=2)
    add_body_text(doc, "Confidence range (68%): ±1 exam-level SEM (±38) around the displayed scaled score — the band within which the resident's true ability falls with 68% probability.")
    add_body_text(doc, "Score band labels (STRONG / ON TRACK / AT RISK / CRITICAL RISK): derived from pass probability estimates using the resident's scaled score position relative to MPS (380) and the national SD. Probabilistic labels, not categorical determinations.")
    add_body_text(doc, "FMCE probability: estimated from the scaled score relative to the FMCE signal threshold (440) using the same normal approximation. A training-exam proxy — does not reflect the FMCE's actual scoring model.")
    add_limitation_callout(doc, "Estimated scaled scores carry ±10–15 point error. Percentile and FMCE probability estimates inherit this error. When official score report PDFs are provided, all displayed values are authoritative ABFM outputs.")

    # =========================================================================
    # SECTION 2: EXAM SUMMARY
    # =========================================================================
    add_section_header(doc, "Section 2 — Exam Summary", level=1)
    add_body_text(doc, "Personal mean: the mean of the resident's per-category rates across all blueprint dimensions. Used as the relative baseline for Tier 2 relative weakness analysis.")
    add_body_text(doc, "Weak areas list: union of two criteria — (1) categories classified as relative weakness by Tier 2 analysis (below personal mean by a statistically meaningful margin), AND (2) categories with absolute rate below 70%. Either criterion qualifies. This list drives practice question targeting and reading list selection.")
    add_body_text(doc, "Easy miss count: questions where item difficulty score >= 700 AND the resident answered incorrectly. Item difficulty (0–999 scale) is ABFM-published from the annual ITE critique — a score of 700 means 70% of national examinees answered correctly.")
    add_body_text(doc, "MPS gap estimate: approximate number of additional correct answers needed to reach the minimum passing score. Linear approximation of the raw-to-scaled relationship, accurate in the 330–480 range.")
    add_limitation_callout(doc, "The MPS gap estimate assumes a linear raw-to-scaled conversion — the actual ABFM conversion uses item response theory and is non-linear. Use as a directional estimate, not a precise target.")

    # =========================================================================
    # SECTION 3B: YEAR-OVER-YEAR PROGRESS
    # =========================================================================
    add_section_header(doc, "Section 3b — Year-Over-Year Progress", level=1)
    add_body_text(doc, "This section appears only when a prior-year analysis is available for the same resident. Both years must be from the same resident to produce a valid comparison.")

    add_section_header(doc, "Trajectory computation", level=2)
    add_body_text(doc, "Scaled delta: current scaled score minus prior scaled score when both are official. Falls back to raw percentage delta if either year used an estimated scaled score.")
    add_body_text(doc, "Trajectory categories per weak-area dimension: Closed = dimension was weak last year but not this year. Persistent = weak in both years. New = weak this year only.")
    add_body_text(doc, "Blueprint delta and body system delta: percentage-point differences (current rate minus prior rate) per dimension. These populate the delta column in Section 4 tables.")
    add_limitation_callout(doc, "Year-over-year comparison is confounded by exam content variation — the ITE is not the same exam each year. A persistent gap may reflect consistent knowledge deficit OR consistent exposure to topics the resident has not yet mastered. Interpret trajectory with clinical context.")

    # =========================================================================
    # SECTION 4: BLUEPRINT & BODY SYSTEM
    # =========================================================================
    add_section_header(doc, "Section 4 — Blueprint & Body System Performance", level=1)
    add_section_header(doc, "Blueprint data", level=2)
    add_data_source_callout(doc, "The official ABFM score report PDF includes a 'Performance by Blueprint Category' table. The extraction stage reads category name, number correct, number of items, and rate from that table. This is authoritative ABFM data. SEM values per category come from ABFM's published annual reference data — ABFM-published measurement error for blueprint sub-scores.")

    add_section_header(doc, "Body system — two-tier provenance", level=2)
    add_data_source_callout(doc, "ABFM-Reported (navy label): parsed from the 'Performance by Body System' section of the official score report PDF. Most reliable — directly from ABFM. Database-Derived (blue label): when ABFM body system data is absent or sparse, a supplemental analysis queries the question database for each missed question's body system classification and aggregates by system. The body system classifications were validated against ABFM's published distributions using a two-stage classifier in April 2026 — 376 assignments were corrected in that validation run.")
    add_limitation_callout(doc, "Database-Derived body systems inherit any remaining errors in the question-level body system assignments. Some mis-assignments may exist at domain boundaries (e.g., a mental health question about a physical complaint). If the two sources diverge, acknowledge the discrepancy rather than treating either as definitive. ABFM-Reported covers only the body systems ABFM chose to report; Database-Derived covers all systems in the question database.")

    # =========================================================================
    # SECTION 6: DIFFICULTY PROFILE
    # =========================================================================
    add_section_header(doc, "Section 6 — Difficulty Profile", level=1)
    add_data_source_callout(doc, "Item difficulty scores are ABFM-published in the annual ITE critique document. The difficulty statistic is the P-value x 1000 — the proportion of national examinees who answered correctly, on a 0–999 scale. These are stored in the question database and matched to the resident's missed items. Tier thresholds: Easy Miss >=700, Mid-Range 300–699, Hard Miss <300.")
    add_limitation_callout(doc, "Difficulty scores from one year do not transfer exactly to another year's exam. Item difficulty fluctuates with the examinee pool and question revision. For questions without published critique data (primarily 2018–2019), difficulty is unavailable and those items are excluded from the difficulty profile.")

    # =========================================================================
    # SECTION 7: LOW-HANGING FRUIT
    # =========================================================================
    add_section_header(doc, "Section 7 — Low-Hanging Fruit", level=1)
    add_section_header(doc, "Priority score formula", level=2)
    add_body_text(doc, "Improvable Items (per dimension): count of missed items where item difficulty >= 500 AND dimension matches. The 500 threshold means more than half of national examinees answered correctly — a proxy for learnability.")
    add_body_text(doc, "Priority Score formula: improvable_items x exam_weight_pct. No normalization or log transformation. A category worth 35% with 3 improvable items scores 105; a category worth 5% with 5 items scores 25. The explicit intent is to weight toward high-exam-weight categories.")
    add_body_text(doc, "Exam weights: ABFM-published blueprint allocations from the annual reference data. Body System table shows Improvable Items only — no Priority Score column because body systems have no ABFM-published exam weight.")

    add_section_header(doc, "Cross-tab (Section 7b)", level=2)
    add_body_text(doc, "Improvable items computed over blueprint x body system intersections. The intersection rate = resident's rate on questions where both the blueprint category AND the body system match the specified pair. Priority score = improvable_items x combined exam weight. Top 8 intersections displayed.")
    add_limitation_callout(doc, "The priority score model is a deliberate simplification. It does not account for diminishing returns, the resident's prior familiarity level, or study resource availability. Use it as a starting sequence, not a precise optimization.")

    # =========================================================================
    # SECTION 8: CONCEPT FINGERPRINT
    # =========================================================================
    add_section_header(doc, "Section 8 — Concept Fingerprint", level=1)
    add_data_source_callout(doc, "Each question in the database has a set of clinical concept tags — an array of drugs, diagnoses, and guideline references associated with that question. These were generated via a clinical concept extraction process and normalized through a clinical synonym map (151 entries) to collapse variants (e.g., T2DM -> type 2 diabetes). The report displays drug concepts only. Diagnosis tags run internally for scoring purposes but are not shown to residents or advisors.")
    add_body_text(doc, "Drug name normalization is more reliable than diagnosis normalization at single-exam sample sizes (191 items) — drug names follow standardized pharmacology conventions; diagnosis labels have too many clinical variants to normalize cleanly. This is why the report surfaces drug concepts rather than diagnosis clusters.")
    add_body_text(doc, "Year-over-year badges compare concept names to the prior-year drug list: matches = Persistent, new entries = New. Up to 10 concepts stored, top 5 displayed.")
    add_limitation_callout(doc, "Concept fingerprint reliability scales with the number of missed items. Residents with fewer than 30 misses will have sparse frequency tables. A drug appearing 2x in 20 misses carries more significance than 2x in 100 misses. Interpret low-miss-count fingerprints proportionally.")

    # =========================================================================
    # SECTION 9: ICD-10 WEAKNESS MAP
    # =========================================================================
    add_section_header(doc, "Section 9 — ICD-10 Weakness Map", level=1)
    add_section_header(doc, "Data flow chain", level=2)

    icd10_table_data = [
        ("Step 1", "Resident's missed question IDs"),
        ("Step 2", "Question-to-article reference table — links each question to one or more clinical guideline articles"),
        ("Step 3", "Article ICD-10 tag table — ICD-10 codes and descriptions assigned to each article"),
        ("Step 4", "Aggregation: miss count per ICD-10 code = number of distinct missed questions whose linked articles include that code"),
    ]
    add_data_table(doc, icd10_table_data)

    add_data_source_callout(doc, "Each of the approximately 2,000 articles in the database was tagged with clinically relevant ICD-10 codes via a clinical coding process. Average 2–3 codes per article. Tagging completed April 2026. ICD-10 chapter mapping is derived from code prefix. Top 15 codes displayed.")
    add_limitation_callout(doc, "This section reflects article-level ICD-10 tags, not question-level tags directly. If a question's linked article covers type 2 diabetes (E11), the question inherits E11 even if the specific question tested metformin dosing rather than diabetes diagnosis — an intentional approximation for pattern detection, not precise question-level coding. Also: question-to-article linkage coverage is 100% for ITE 2018–2023, 90% for 2024, and 83.5% for 2025 — questions without article linkages will not contribute to the ICD-10 map.")

    # =========================================================================
    # SECTION 10: HIGH-YIELD READING LIST
    # =========================================================================
    add_section_header(doc, "Section 10 — High-Yield Reading List", level=1)
    add_section_header(doc, "Two-tier selection logic", level=2)
    add_body_text(doc, "Tier 1 — Personalized (Targeted to Your Exam): articles are selected from the pool of articles linked to questions the resident answered incorrectly. Articles with two or more links to the resident's weak-area missed questions are selected first (Strong Tier 1); articles with one link are used as overflow. Target 5 articles, cap 7.")
    add_body_text(doc, "Tier 2 — General (High-Yield for All Residents): articles not selected in Tier 1. Strong Tier 2 requires high citation frequency across multiple exam years; Overflow Tier 2 requires moderate citation frequency. Target 5 articles, cap 8. These represent cornerstone FM literature — articles that have appeared repeatedly across exam years.")

    add_section_header(doc, "Supporting data", level=2)
    add_data_source_callout(doc, "Article selection draws from three sources: the article database (~2,000 articles), the question-to-article reference table (linking exam questions to their source guidelines), and article currency data populated via a PubMed lookup process. Currency statuses: current / updated / check_needed / not_indexed. Linked question IDs are attached to each article in the reading list — up to 10 stored, top 5 displayed.")
    add_limitation_callout(doc, "Tier 1 personalization quality depends on question-to-article coverage. For 2024–2025 exams (83–90% coverage), some missed questions may not have linked articles, reducing the precision of personalized recommendations. Currency status reflects a one-time lookup, not live monitoring — verify independently for any article marked 'updated' or 'check_needed'.")

    # =========================================================================
    # PART B: PRACTICE QUESTION SELECTION
    # =========================================================================
    add_section_header(doc, "Part B — Practice Question Selection", level=1)
    add_section_header(doc, "Three-tier matching cascade", level=2)

    pq_table_data = [
        ("Tier 1 — Direct Match", "Questions sharing the same source article as the resident's actual exam misses. Highest precision — these practice questions are clinically adjacent to the resident's specific knowledge gaps. Shown in GREEN."),
        ("Tier 2 — ICD-10 Sibling", "Questions not matched in Tier 1 whose ICD-10 tags overlap with the resident's weak-area clinical profile. Clinical relatives of the missed items, drawn from the same disease territory. Shown in BLUE."),
        ("Tier 3 — Vector Match", "Questions not matched in Tier 1 or 2 where the question's semantic embedding is similar to the resident's weak-area centroid. Broadest net — useful when Tier 1 and 2 pools are small. Shown in GRAY."),
    ]
    add_data_table(doc, pq_table_data)

    add_body_text(doc, "Relevance score: composite of tier membership + concept tag overlap + ICD-10 overlap + vector similarity. Questions sorted descending by relevance score. Source pool: all ITE questions from prior exam years plus all AAFP board review questions. A question is never drawn from the same exam being analyzed.")
    add_limitation_callout(doc, "Tier 3 vector matching can surface questions that are semantically similar but not topically identical — they may test adjacent knowledge rather than the exact knowledge gap. If the practice set feels off-topic, the likely culprit is the Tier 3 (Vector Match) questions.")

    # =========================================================================
    # APPENDIX: MISSED EXAM ITEMS
    # =========================================================================
    add_section_header(doc, "Appendix — Missed Exam Items", level=1)
    add_body_text(doc, "All missed items are stored in the analysis output and rendered in the appendix as a reference table for question cross-referencing. No new analysis occurs in this section.")
    add_body_text(doc, "Scoring note: The ITE contains 200 numbered items, but the scored count is typically 191. The discrepancy reflects ABFM's practice of embedding experimental (unscored) items. These appear in the question numbering but are excluded from the scaled score calculation. The appendix note explains this to the resident.")

    # =========================================================================
    # FACULTY ADVISING FRAMEWORK
    # =========================================================================
    add_section_header(doc, "Faculty Advising Framework", level=1)
    add_body_text(doc, "This section synthesizes the report into advising decision points. It is not tied to a specific report section — it is guidance for the advising conversation.")

    add_section_header(doc, "When to act urgently", level=2)
    add_bullet(doc, "Scaled score below MPS (380) AND persistent gaps across 2+ years AND easy miss count >8 — advising intervention before next exam cycle, with a structured study plan and check-in schedule.")
    add_bullet(doc, "PGY-3 below PGY-1 national mean — trajectory is moving in the wrong direction at the wrong time.")

    add_section_header(doc, "When to reframe study strategy (not add volume)", level=2)
    add_bullet(doc, "Same weak category across 2+ exam years — the issue is method, not time investment. Ask what resource they have been using and whether it has been effective.")
    add_bullet(doc, "High easy miss count despite reported studying — confirms knowledge exposure but not retention. Recommend active recall over passive review.")

    add_section_header(doc, "Calibrated confidence in the data", level=2)
    add_bullet(doc, "Never advise based solely on a sub-score with SEM >12% without corroboration from another section of the report.")
    add_bullet(doc, "Database-Derived body system scores are supplementary signals — corroborate with the ABFM-Reported table before making high-stakes advising decisions.")
    add_bullet(doc, "Estimated scaled scores (labeled 'est.') carry ±10–15 point error — treat them as directional, not definitive.")

    add_section_header(doc, "Longitudinal tracking", level=2)
    add_body_text(doc, "The most valuable use of this report is not a single exam — it is the year-over-year trajectory. A resident moving from 355 to 390 to 420 over three years is on a better trajectory than one stuck at 400 despite additional studying. Use the Year-Over-Year section as your primary advising anchor for returning residents.")

    add_section_header(doc, "What the report cannot tell you", level=2)
    add_body_text(doc, "Clinical performance in continuity clinic, patient care skills, procedural competence, or professional development. The ITE measures knowledge only. Advise holistically — this report is one data point in a comprehensive resident assessment.")

    add_divider(doc)

    # FOOTER & SAVE
    add_page_number_footer(doc.sections[0])
    doc.save(str(OUTPUT_PATH))

    print(f"✓ Faculty advisor guide generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_faculty_guide()
