"""
aafp_exam_builder.py
====================
Builds AAFP BRQ exam documents (exam + explanations) for ITE-linked questions.

Outputs 4 DOCX files:
  - AAFP_ITE_Exam_1.docx        300 questions (random)
  - AAFP_ITE_Exam_2.docx        300 questions (random, from remainder)
  - AAFP_ITE_Exam_3.docx         43 questions (remainder)
  - AAFP_ITE_Exam_Explanations.docx  all 643, ordered by exam then Q#

Usage:
  python aafp_exam_builder.py [--seed N] [--output-dir PATH]

Flags:
  --seed N         Random seed for reproducible splits (default: 42)
  --output-dir     Output directory (default: project root / aafp_exam_docs)

Locked rules:
  - from word_doc_defaults import *  (St. Luke's palette, Aptos font, US Letter)
  - SOURCE DATA PROTECTED — read-only DB access
  - SCRIPT_DIR / PROJECT_ROOT dynamic path pattern
"""

import argparse
import json
import random
import sqlite3
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Dynamic paths ─────────────────────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
DB_PATH      = PROJECT_ROOT / "00_database" / "db" / "ite_intelligence.db"

sys.path.insert(0, str(SCRIPT_DIR))
from word_doc_defaults import (
    new_document, apply_page_setup, set_default_font,
    add_title, add_subtitle, add_section_header, add_body_text,
    add_answer_choice, add_correct_answer_label, add_divider,
    add_page_number_footer, add_shading, add_left_border, add_bottom_border,
    set_paragraph_format, sanitize,
    NAVY, GOLD, BLUE, LIGHT_BLUE, DARK_TEXT, MED_GRAY,
    RGB_NAVY, RGB_GOLD, RGB_BLUE, RGB_DARK_TEXT, RGB_GRAY, RGB_GREEN,
    DEFAULT_FONT, FONT_TITLE, FONT_SUBTITLE, FONT_HEADING, FONT_BODY,
    FONT_SMALL, FONT_TINY,
)

TODAY = date.today().strftime("%B %d, %Y")

# ── Config ────────────────────────────────────────────────────────────────────
EXAM_SIZES   = [300, 300, 43]   # must sum to 643
FOOTER_STUB  = "AAFP BRQ  |  St. Luke's Family Medicine Residency"


# ══════════════════════════════════════════════════════════════════════════════
#  DATA LAYER
# ══════════════════════════════════════════════════════════════════════════════

def load_questions():
    """
    Return list of dicts for all AAFP questions linked to ≥1 ITE article,
    deduplicated to one row per aafp_qid.
    """
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    cur.execute("""
        SELECT DISTINCT
            q.aafp_qid,
            q.stem,
            q.choices,
            q.correct_letter,
            q.correct_text,
            q.explanation,
            q.blueprint,
            q.body_system
        FROM aafp_questions q
        JOIN aafp_qid_art_xref x ON q.aafp_qid = x.aafp_qid
        ORDER BY q.aafp_qid
    """)
    rows = cur.fetchall()
    conn.close()

    questions = []
    for row in rows:
        choices_raw = row["choices"] or "[]"
        try:
            choices = json.loads(choices_raw)
        except json.JSONDecodeError:
            choices = []
        questions.append({
            "aafp_qid":      row["aafp_qid"],
            "stem":          sanitize(row["stem"] or ""),
            "choices":       choices,
            "correct_letter": row["correct_letter"] or "",
            "correct_text":  sanitize(row["correct_text"] or ""),
            "explanation":   sanitize(row["explanation"] or ""),
            "blueprint":     row["blueprint"] or "",
            "body_system":   row["body_system"] or "",
        })
    return questions


def split_questions(questions, sizes, seed=42):
    """Randomly shuffle and split into groups of sizes[0], sizes[1], ..."""
    rng = random.Random(seed)
    pool = questions[:]
    rng.shuffle(pool)
    groups = []
    idx = 0
    for size in sizes:
        groups.append(pool[idx : idx + size])
        idx += size
    return groups


# ══════════════════════════════════════════════════════════════════════════════
#  SHARED HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def set_footer(doc, stub=FOOTER_STUB):
    """Apply footer to every section in the document."""
    for section in doc.sections:
        _write_footer_para(section, stub)


def _write_footer_para(section, stub):
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Right-aligned tab stop for page number
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')
    tabs.append(tab)
    pPr.append(tabs)

    left_run = p.add_run(stub)
    left_run.font.name      = DEFAULT_FONT
    left_run.font.size      = Pt(FONT_TINY)
    left_run.font.color.rgb = RGB_GRAY

    p.add_run("\t").font.size = Pt(FONT_TINY)

    page_run = p.add_run("Page ")
    page_run.font.name      = DEFAULT_FONT
    page_run.font.size      = Pt(FONT_TINY)
    page_run.font.color.rgb = RGB_GRAY

    fld = p.add_run()
    fld.font.size      = Pt(FONT_TINY)
    fld.font.color.rgb = RGB_GRAY
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
    ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'), 'preserve'); ins.text = 'PAGE'
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end')
    fld._r.extend([fc1, ins, fc2])


def add_cover(doc, title, subtitle, instruction_lines):
    """Add a cover page with title, subtitle, and instruction block."""
    # spacer at top
    for _ in range(6):
        doc.add_paragraph()

    add_title(doc, title)
    add_subtitle(doc, subtitle)

    # thin spacer
    doc.add_paragraph()

    # instructions box
    for line in instruction_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name  = DEFAULT_FONT
        run.font.size  = Pt(FONT_SMALL)
        run.font.color.rgb = RGB_GRAY

    # page break after cover
    doc.add_paragraph().add_run().add_break(
        __import__('docx.enum.text', fromlist=['WD_BREAK_TYPE']).WD_BREAK_TYPE.PAGE
    )


def add_question_block(doc, number, q):
    """Render one question: numbered stem + lettered choices."""
    # Question stem
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    set_paragraph_format(p, keep_with_next=True)

    # Number (bold navy)
    num_run = p.add_run(f"{number}.  ")
    num_run.font.name      = DEFAULT_FONT
    num_run.font.size      = Pt(FONT_BODY)
    num_run.font.bold      = True
    num_run.font.color.rgb = RGB_NAVY

    # Stem text
    stem_run = p.add_run(q["stem"])
    stem_run.font.name      = DEFAULT_FONT
    stem_run.font.size      = Pt(FONT_BODY)
    stem_run.font.color.rgb = RGB_DARK_TEXT

    # Choices
    for choice in q["choices"]:
        letter = choice.get("letter", "?")
        text   = sanitize(choice.get("text", ""))
        add_answer_choice(doc, letter, text, indent=0.4)

    # Thin divider between questions
    add_divider(doc)


def add_answer_key_section(doc, group, start_num=1):
    """
    Render the Answer Key as a 3-column table: Q# | Answer | Correct Text.
    Appended on a new page at the end of the exam document.
    """
    # Page break
    doc.add_paragraph().add_run().add_break(
        __import__('docx.enum.text', fromlist=['WD_BREAK_TYPE']).WD_BREAK_TYPE.PAGE
    )

    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_hdr.paragraph_format.space_after = Pt(6)
    run = p_hdr.add_run("Answer Key")
    run.font.name      = DEFAULT_FONT
    run.font.size      = Pt(16)
    run.font.bold      = True
    run.font.color.rgb = RGB_NAVY
    add_bottom_border(p_hdr, color=NAVY, size=8)

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt as _Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WDA

    # Table: col widths 1080 + 900 + 7380 = 9360 DXA
    COL_NUM  = 1080
    COL_ANS  = 900
    COL_TEXT = 7380
    TOTAL    = COL_NUM + COL_ANS + COL_TEXT

    border = {"val": "single", "sz": "4", "color": "CCCCCC"}
    def cell_border():
        tc_pr = OxmlElement('w:tcPr')
        tc_borders = OxmlElement('w:tcBorders')
        for side in ['top', 'bottom', 'left', 'right']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),   'single')
            b.set(qn('w:sz'),    '4')
            b.set(qn('w:color'), 'CCCCCC')
            tc_borders.append(b)
        tc_pr.append(tc_borders)
        return tc_pr

    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    from docx.oxml import OxmlElement

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(TOTAL))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    def set_cell_width(cell, width_dxa):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    def shade_cell(cell, fill):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill)
        tcPr.append(shd)

    def style_cell_para(cell, text, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, size=FONT_SMALL):
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.left_indent  = Inches(0.05)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text)
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        return p

    # Header row
    hdr_row = table.rows[0]
    headers = [("Q#", COL_NUM), ("Ans", COL_ANS), ("Correct Answer", COL_TEXT)]
    for i, (hdr_text, width) in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_width(cell, width)
        shade_cell(cell, "1B3564")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(hdr_text)
        run.font.name      = DEFAULT_FONT
        run.font.size      = Pt(FONT_SMALL)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, q in enumerate(group):
        qnum = start_num + i
        row = table.add_row()
        fill = "EFF3FA" if i % 2 == 0 else "FFFFFF"

        # Q#
        c0 = row.cells[0]
        set_cell_width(c0, COL_NUM)
        shade_cell(c0, fill)
        style_cell_para(c0, str(qnum), bold=True, color=RGB_NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Answer letter
        c1 = row.cells[1]
        set_cell_width(c1, COL_ANS)
        shade_cell(c1, fill)
        style_cell_para(c1, q["correct_letter"], bold=True, color=RGB_GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Correct text
        c2 = row.cells[2]
        set_cell_width(c2, COL_TEXT)
        shade_cell(c2, fill)
        style_cell_para(c2, q["correct_text"], bold=False, color=RGB_DARK_TEXT)

    doc.add_paragraph()  # trailing space


# ══════════════════════════════════════════════════════════════════════════════
#  EXAM DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_exam_doc(exam_num, group, output_path):
    """Build one exam DOCX: cover → questions → answer key."""
    n = len(group)
    print(f"  Building Exam {exam_num}: {n} questions → {output_path.name}")

    doc = new_document()

    instructions = [
        f"Questions: {n}     |     Estimated time: {n * 1.2:.0f} minutes",
        "Select the single best answer for each question.",
        f"Generated: {TODAY}     |     Source: AAFP BRQ (ITE-Linked)"
    ]
    add_cover(
        doc,
        title=f"AAFP Board Review Questions — Exam {exam_num}",
        subtitle="ITE-Linked Question Bank  |  St. Luke's Family Medicine Residency",
        instruction_lines=instructions,
    )

    # Questions
    for i, q in enumerate(group):
        add_question_block(doc, i + 1, q)

    # Answer Key
    add_answer_key_section(doc, group, start_num=1)

    set_footer(doc)
    doc.save(str(output_path))
    print(f"    ✓ Saved: {output_path}")


# ══════════════════════════════════════════════════════════════════════════════
#  EXPLANATIONS DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_explanations_doc(groups, output_path):
    """
    Build the Explanations DOCX.
    Sections: Exam 1 Explanations, Exam 2 Explanations, Exam 3 Explanations.
    Each entry: Q# header → correct answer → explanation text.
    """
    total = sum(len(g) for g in groups)
    print(f"  Building Explanations doc: {total} questions → {output_path.name}")

    doc = new_document()

    # Cover
    instructions = [
        f"Covers Exams 1, 2, and 3  |  {total} total questions",
        "Q# numbering matches corresponding exam document.",
        f"Generated: {TODAY}     |     Source: AAFP BRQ (ITE-Linked)"
    ]
    add_cover(
        doc,
        title="AAFP Board Review Questions — Explanations",
        subtitle="ITE-Linked Question Bank  |  St. Luke's Family Medicine Residency",
        instruction_lines=instructions,
    )

    for exam_num, group in enumerate(groups, start=1):
        # Exam section header — new page (except first, already on new page)
        if exam_num > 1:
            doc.add_paragraph().add_run().add_break(
                __import__('docx.enum.text', fromlist=['WD_BREAK_TYPE']).WD_BREAK_TYPE.PAGE
            )

        # Large section header
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(6)
        p_sec.paragraph_format.space_after  = Pt(10)
        run = p_sec.add_run(f"Exam {exam_num} — Explanations  ({len(group)} questions)")
        run.font.name      = DEFAULT_FONT
        run.font.size      = Pt(16)
        run.font.bold      = True
        run.font.color.rgb = RGB_NAVY
        add_bottom_border(p_sec, color=NAVY, size=8)

        for i, q in enumerate(group):
            qnum = i + 1

            # Q# sub-header with gold left bar
            hdr_p = doc.add_paragraph()
            hdr_p.paragraph_format.space_before = Pt(14)
            hdr_p.paragraph_format.space_after  = Pt(3)
            add_left_border(hdr_p, color=GOLD, size=24)
            add_shading(hdr_p, "EFF3FA")
            set_paragraph_format(hdr_p, keep_with_next=True)

            num_run = hdr_p.add_run(f"Q{qnum}  ")
            num_run.font.name      = DEFAULT_FONT
            num_run.font.size      = Pt(FONT_HEADING)
            num_run.font.bold      = True
            num_run.font.color.rgb = RGB_NAVY

            # Blueprint category (gray)
            if q["blueprint"]:
                cat_run = hdr_p.add_run(f"  [{q['blueprint']}]")
                cat_run.font.name      = DEFAULT_FONT
                cat_run.font.size      = Pt(FONT_SMALL)
                cat_run.font.bold      = False
                cat_run.font.color.rgb = RGB_GRAY

            # Stem (lighter, for context)
            stem_p = doc.add_paragraph()
            stem_p.paragraph_format.left_indent = Inches(0.25)
            stem_p.paragraph_format.space_after = Pt(4)
            set_paragraph_format(stem_p, keep_with_next=True)
            stem_run = stem_p.add_run(q["stem"])
            stem_run.font.name      = DEFAULT_FONT
            stem_run.font.size      = Pt(FONT_SMALL)
            stem_run.font.color.rgb = RGB_GRAY
            stem_run.font.italic    = True

            # Correct answer label (green)
            ca_p = doc.add_paragraph()
            ca_p.paragraph_format.left_indent  = Inches(0.25)
            ca_p.paragraph_format.space_before = Pt(4)
            ca_p.paragraph_format.space_after  = Pt(2)
            set_paragraph_format(ca_p, keep_with_next=True)
            ca_run = ca_p.add_run(f"CORRECT ANSWER: {q['correct_letter']}  —  {q['correct_text']}")
            ca_run.font.name      = DEFAULT_FONT
            ca_run.font.size      = Pt(FONT_BODY)
            ca_run.font.bold      = True
            ca_run.font.color.rgb = RGB_GREEN

            # Explanation body
            exp_p = doc.add_paragraph()
            exp_p.paragraph_format.left_indent = Inches(0.25)
            exp_p.paragraph_format.space_after  = Pt(6)
            exp_run = exp_p.add_run(q["explanation"])
            exp_run.font.name      = DEFAULT_FONT
            exp_run.font.size      = Pt(FONT_SMALL)
            exp_run.font.color.rgb = RGB_DARK_TEXT

            add_divider(doc)

    set_footer(doc)
    doc.save(str(output_path))
    print(f"    ✓ Saved: {output_path}")


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Build AAFP BRQ exam documents.")
    parser.add_argument("--seed",       type=int,  default=42,   help="Random seed (default: 42)")
    parser.add_argument("--output-dir", type=str,  default=None, help="Output directory")
    args = parser.parse_args()

    # Output dir
    if args.output_dir:
        out_dir = Path(args.output_dir)
    else:
        out_dir = PROJECT_ROOT / "aafp_exam_docs"
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n{'='*60}")
    print(f"AAFP Exam Builder")
    print(f"  DB:         {DB_PATH}")
    print(f"  Output dir: {out_dir}")
    print(f"  Seed:       {args.seed}")
    print(f"{'='*60}\n")

    # Load + split
    print("Loading questions from DB...")
    questions = load_questions()
    print(f"  Loaded {len(questions)} ITE-linked AAFP questions")

    assert len(questions) == sum(EXAM_SIZES), (
        f"Expected {sum(EXAM_SIZES)} questions, got {len(questions)}. "
        "Update EXAM_SIZES if DB has changed."
    )

    groups = split_questions(questions, EXAM_SIZES, seed=args.seed)
    for i, g in enumerate(groups):
        print(f"  Exam {i+1}: {len(g)} questions")

    print()

    # Build exam docs
    for exam_num, group in enumerate(groups, start=1):
        out_path = out_dir / f"AAFP_ITE_Exam_{exam_num}.docx"
        build_exam_doc(exam_num, group, out_path)

    print()

    # Build explanations doc
    exp_path = out_dir / "AAFP_ITE_Exam_Explanations.docx"
    build_explanations_doc(groups, exp_path)

    print(f"\n{'='*60}")
    print(f"Done. 4 files written to: {out_dir}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
