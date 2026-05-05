#!/usr/bin/env python3
"""
build_cole_exam_series.py
=========================
Generates 4 practice exam files (A/B/C/D) for Okezia Cole, PGY-1, 2025 ITE.
Each file contains 50 unique questions weighted by weakness profile.
Questions are randomly distributed across the 4 files.

Also cleans the Symbol-font dot-leader encoding artifact (U+F02E → space) from
analysis_v2_2025.json in-place, so re-running the JS builder also produces a
clean single exam file.

Rule 14 compliant: uses word_doc_defaults.py for all DOCX generation.
Rule 4 compliant: dynamic paths only.

Usage:
    cd 03_module.3_analyst/scripts
    python3 build_cole_exam_series.py

Outputs (all in resident_data/ITE_okezia_cole/outputs/):
    analysis_v2_2025.json             -- encoding-cleaned in place
    ITE_2025_Practice_Exam_A_Cole.docx
    ITE_2025_Practice_Exam_B_Cole.docx
    ITE_2025_Practice_Exam_C_Cole.docx
    ITE_2025_Practice_Exam_D_Cole.docx

After running, also run to regenerate the fixed single exam:
    node ite_report_builder_v2.js \\
        "../resident_data/ITE_okezia_cole/outputs/analysis_v2_2025.json" \\
        "../resident_data/ITE_okezia_cole/outputs/"
"""

from __future__ import annotations

import json
import random
import re
import sys
from collections import Counter
from pathlib import Path

# ── Path setup (Rule 4: dynamic paths only) ───────────────────────────────────
SCRIPT_DIR   = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent

DB_PATH       = PROJECT_ROOT / "00_database" / "db" / "ite_intelligence.db"
RESIDENT_DIR  = PROJECT_ROOT / "03_module.3_analyst" / "resident_data" / "ITE_okezia_cole"
ANALYSIS_JSON = RESIDENT_DIR / "outputs" / "analysis_v2_2025.json"
OUTPUT_DIR    = RESIDENT_DIR / "outputs"

sys.path.insert(0, str(SCRIPT_DIR))

# Rule 14: import word_doc_defaults
from word_doc_defaults import (
    new_document, add_title, add_subtitle,
    add_bottom_border, add_left_border, add_shading, set_paragraph_format,
    DEFAULT_FONT, RGB_NAVY, RGB_BLUE, RGB_GRAY, RGB_GOLD,
    NAVY, BLUE, GOLD,
    FONT_TITLE, FONT_SUBTITLE, FONT_HEADING, FONT_BODY, FONT_SMALL, FONT_TINY,
)
from ite_analyzer_v3 import match_practice_questions_v3

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Encoding fix ──────────────────────────────────────────────────────────────
# U+F02E = Symbol-font private-use period (dot-leader in original ITE PDF).
# Renders as garbage (ï€®) in Word without the Symbol font embedded.
_DOTLEADER = re.compile(r'[]+')

def clean_text(text: str | None) -> str:
    """Replace Symbol-font dot-leader chars and normalize Unicode punctuation."""
    if not text:
        return text or ""
    text = _DOTLEADER.sub(" ", text)
    text = text.replace("“", '"').replace("”", '"')
    text = text.replace("‘", "'").replace("’", "'")
    return text

def clean_choices(raw) -> dict:
    """Normalize choices to {letter: text} dict regardless of source format."""
    if isinstance(raw, str):
        try:
            parsed = json.loads(raw)
        except Exception:
            return {}
        if isinstance(parsed, list):
            return {c["letter"]: clean_text(c.get("text", "")) for c in parsed if "letter" in c}
        if isinstance(parsed, dict):
            return {k: clean_text(v) for k, v in parsed.items()}
        return {}
    if isinstance(raw, list):
        return {c["letter"]: clean_text(c.get("text", "")) for c in raw if "letter" in c}
    if isinstance(raw, dict):
        return {k: clean_text(v) for k, v in raw.items()}
    return {}


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 1 — Load and clean analysis JSON in-place
# ═══════════════════════════════════════════════════════════════════════════════
print("=" * 60)
print("STEP 1: Clean encoding in analysis_v2_2025.json")
print("=" * 60)

with open(ANALYSIS_JSON, encoding="utf-8") as f:
    analysis = json.load(f)

mid = analysis.get("missed_items_detail", [])
cleaned_count = 0
for m in mid:
    if "" in (m.get("question_text") or ""):
        m["question_text"] = clean_text(m["question_text"])
        cleaned_count += 1
    if m.get("explanation"):
        m["explanation"] = clean_text(m["explanation"])

for q in analysis.get("practice_questions", []):
    if q.get("question_text"):
        q["question_text"] = clean_text(q["question_text"])
    if q.get("explanation"):
        q["explanation"] = clean_text(q["explanation"])

with open(ANALYSIS_JSON, "w", encoding="utf-8") as f:
    json.dump(analysis, f, indent=2, ensure_ascii=False)

print(f"  ✓ {cleaned_count} question(s) with dot-leader artifacts cleaned")
print(f"  ✓ JSON saved — single exam will be clean on next JS builder run")


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 2 — Pull 200 weakness-weighted questions
# ═══════════════════════════════════════════════════════════════════════════════
print()
print("=" * 60)
print("STEP 2: Pull 200 questions from DB")
print("=" * 60)

# Reconstruct items list from missed_items_detail
# match_practice_questions_v3 uses items to identify which QIDs belong to
# each weak dimension — we only have missed items (all correct=False)
items = [
    {
        "item":        m["item_number"],
        "correct":     False,
        "blueprint":   m.get("blueprint"),
        "body_system": m.get("body_system"),
        "score":       500,   # not used by match_practice_questions_v3
    }
    for m in mid if m.get("item_number")
]

qid_map = {
    m["item_number"]: m["qid"]
    for m in mid if m.get("item_number") and m.get("qid")
}

# ICD-10 profile (invisible scoring signal per Rule 13)
icd10_clusters = analysis.get("icd10_weakness_map", {}).get("icd10_clusters", [])
icd10_profile  = {
    c["code"]: c["miss_count"]
    for c in icd10_clusters
    if c.get("code") and c.get("miss_count")
}

priorities  = analysis.get("yield_priorities", [])
concepts    = analysis.get("concept_clustering", {})
exam_year   = str(analysis.get("exam_year", "2025"))

print(f"  Items (missed): {len(items)}")
print(f"  Priorities (weak dims): {len(priorities)}")
print(f"  ICD-10 profile codes: {len(icd10_profile)}")
print(f"  DB: {DB_PATH}")
print()

perf = analysis.get("performance", {})

all_200 = match_practice_questions_v3(
    perf=perf,
    priorities=priorities,
    qid_map=qid_map,
    items=items,
    db_path=str(DB_PATH),
    target_count=200,
    concepts=concepts,
    current_exam_year=exam_year,
    icd10_profile=icd10_profile,
)

print(f"\n  Questions returned: {len(all_200)}")
tier_counts = Counter(q.get("match_tier") for q in all_200)
print(f"  Match tiers: " + "  ".join(f"T{k}:{v}" for k, v in sorted(tier_counts.items())))

bank_counts = Counter(q.get("source_bank") for q in all_200)
print(f"  Source banks: " + "  ".join(f"{k}:{v}" for k, v in sorted(bank_counts.items())))

bp_counts = Counter(q.get("blueprint") for q in all_200)
print(f"\n  Blueprint distribution:")
for bp, cnt in sorted(bp_counts.items(), key=lambda x: -x[1]):
    print(f"    {bp}: {cnt}")

# Apply encoding fix to all returned questions
for q in all_200:
    q["question_text"] = clean_text(q.get("question_text", ""))
    q["explanation"]   = clean_text(q.get("explanation", ""))
    q["choices"]       = clean_choices(q.get("choices"))

# ── Inject original report questions ──────────────────────────────────────────
# The 20 practice questions already in the analysis report should appear in the
# exam series too.  Find any that weren't selected in the top-200, inject them,
# and drop an equal number from the BOTTOM of the ranked list (least fitting).
print()
print("  Merging original report questions into pool…")

all_200_qids   = {q.get("qid") for q in all_200 if q.get("qid")}
original_qs    = analysis.get("practice_questions", [])
missing_from_200 = [
    q for q in original_qs
    if q.get("qid") and q.get("qid") not in all_200_qids
]

n_inject = len(missing_from_200)
if n_inject == 0:
    print(f"  ✓ All {len(original_qs)} original questions already in pool — no change")
    final_pool = all_200
else:
    print(f"  ✓ {n_inject} original question(s) not in ranked pool")
    print(f"    Injecting {n_inject} | Clipping {n_inject} least-fitting from bottom of ranked list")
    # Apply encoding fix to injected questions
    for q in missing_from_200:
        q["question_text"] = clean_text(q.get("question_text", ""))
        q["explanation"]   = clean_text(q.get("explanation", ""))
        q["choices"]       = clean_choices(q.get("choices"))
    # Keep top (200 - n_inject) from ranked list + inject originals
    final_pool = all_200[:-n_inject] + missing_from_200

print(f"  Final pool: {len(final_pool)} questions")


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 3 — Randomly distribute into 4 groups of 50
# ═══════════════════════════════════════════════════════════════════════════════
print()
print("=" * 60)
print("STEP 3: Distribute questions across 4 exams")
print("=" * 60)

actual_count = len(final_pool)
if actual_count < 200:
    print(f"  ⚠ Only {actual_count} unique questions available (requested 200)")
    per_exam = actual_count // 4
else:
    per_exam = 50

random.seed(42)   # reproducible
shuffled = final_pool[:]
random.shuffle(shuffled)

exams = {
    "A": shuffled[0           : per_exam],
    "B": shuffled[per_exam    : per_exam * 2],
    "C": shuffled[per_exam*2  : per_exam * 3],
    "D": shuffled[per_exam*3  : per_exam * 4],
}

for letter, qs in exams.items():
    t = Counter(q.get("match_tier") for q in qs)
    b = Counter(q.get("source_bank") for q in qs)
    print(f"  Exam {letter}: {len(qs)} questions | "
          f"T1:{t.get(1,0)} T2:{t.get(2,0)} T3:{t.get(3,0)} | "
          f"ITE:{b.get('ITE',0)} AAFP:{b.get('AAFP',0)}")


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 4 — Build DOCX files
# ═══════════════════════════════════════════════════════════════════════════════

def add_page_header(doc: Document, text: str):
    """Add right-aligned running header to all pages."""
    section = doc.sections[0]
    header  = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name       = DEFAULT_FONT
    run.font.size       = Pt(FONT_TINY)
    run.font.color.rgb  = RGB_GRAY


def add_answer_key_table(doc: Document, answer_key: list):
    """Render answer key as a formatted table."""
    col_headers = ["#", "QID", "Answer", "Target Area", "Body System"]
    col_widths  = [Inches(0.4), Inches(1.55), Inches(0.65), Inches(1.85), Inches(1.55)]

    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = "Table Grid"

    # Header row — navy background, white text
    hdr_cells = table.rows[0].cells
    for i, (hdr, w) in enumerate(zip(col_headers, col_widths)):
        hdr_cells[i].width = w
        p   = hdr_cells[i].paragraphs[0]
        run = p.add_run(hdr)
        run.font.name       = DEFAULT_FONT
        run.font.size       = Pt(FONT_SMALL)
        run.font.bold       = True
        run.font.color.rgb  = RGBColor(0xFF, 0xFF, 0xFF)
        # Apply navy cell background
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  NAVY)
        tcPr.append(shd)

    # Data rows — alternating light blue every other row
    for row_idx, entry in enumerate(answer_key):
        cells = table.add_row().cells
        vals  = [
            str(entry["num"]),
            entry.get("qid", ""),
            entry.get("letter", ""),
            entry.get("target", ""),
            entry.get("body", ""),
        ]
        fill = "EBF0F7" if row_idx % 2 == 0 else "FFFFFF"
        for i, (val, w) in enumerate(zip(vals, col_widths)):
            cells[i].width = w
            p   = cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(FONT_SMALL)
            # Alternate row shading
            tcPr = cells[i]._tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tcPr.append(shd)


def build_exam_docx(questions: list, exam_letter: str, output_path: Path):
    """Build a single practice exam DOCX."""
    doc = new_document()  # Rule 14: uses word_doc_defaults page setup + Aptos font

    add_page_header(doc, f"ITE {exam_year} Practice Exam {exam_letter}")

    # ── Cover page ────────────────────────────────────────────────────────
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(72)

    add_title(doc,    f"ITE {exam_year} PRACTICE EXAM {exam_letter}")
    add_subtitle(doc, f"Okezia Cole, D.O.  |  PGY-1  |  St. Luke's Family Medicine Residency")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(6)
    r = info.add_run(
        f"{len(questions)} Questions — Weighted by Performance Priority  |  Answer key at end"
    )
    r.font.name      = DEFAULT_FONT
    r.font.size      = Pt(FONT_SMALL)
    r.font.color.rgb = RGB_GRAY

    doc.add_page_break()

    # ── Questions ─────────────────────────────────────────────────────────
    answer_key = []

    for i, q in enumerate(questions, 1):
        q_text   = q.get("question_text", "")
        choices  = q.get("choices") or {}
        targeting = q.get("targeting", "")
        bank      = q.get("source_label") or q.get("source_bank", "")
        body_sys  = q.get("body_system_merged") or q.get("body_system", "")
        qid       = q.get("qid", "")
        correct_letter = q.get("correct_letter", "")
        correct_text   = q.get("correct_text", "")

        answer_key.append({
            "num":    i,
            "qid":    qid,
            "letter": correct_letter,
            "text":   correct_text,
            "target": targeting,
            "body":   body_sys,
        })

        # "Question N" label
        qh = doc.add_paragraph()
        qh.paragraph_format.space_before = Pt(16)
        qh.paragraph_format.space_after  = Pt(4)
        set_paragraph_format(qh, keep_with_next=True)
        rn = qh.add_run(f"Question {i}")
        rn.font.name      = DEFAULT_FONT
        rn.font.size      = Pt(FONT_HEADING)
        rn.font.bold      = True
        rn.font.color.rgb = RGB_NAVY

        # Question stem
        qs = doc.add_paragraph()
        qs.paragraph_format.space_before = Pt(2)
        qs.paragraph_format.space_after  = Pt(6)
        set_paragraph_format(qs, keep_with_next=True)
        rt = qs.add_run(q_text)
        rt.font.name = DEFAULT_FONT
        rt.font.size = Pt(FONT_BODY)

        # Answer choices
        if isinstance(choices, dict):
            choice_items = sorted(choices.items())
        elif isinstance(choices, list):
            choice_items = [(c.get("letter", ""), c.get("text", "")) for c in choices]
        else:
            choice_items = []

        for letter, choice_text in choice_items:
            cp = doc.add_paragraph()
            cp.paragraph_format.left_indent  = Inches(0.35)
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after  = Pt(2)
            rc = cp.add_run(f"{letter}. {clean_text(choice_text)}")
            rc.font.name = DEFAULT_FONT
            rc.font.size = Pt(FONT_BODY)

        # Thin divider between questions
        if i < len(questions):
            div = doc.add_paragraph()
            div.paragraph_format.space_before = Pt(10)
            div.paragraph_format.space_after  = Pt(0)
            add_bottom_border(div, color="CCCCCC", size=4)

    # ── Answer key page ───────────────────────────────────────────────────
    doc.add_page_break()

    ak_hdr = doc.add_paragraph()
    ak_hdr.paragraph_format.space_before = Pt(0)
    ak_hdr.paragraph_format.space_after  = Pt(12)
    add_left_border(ak_hdr, color=GOLD, size=24)
    add_shading(ak_hdr, "EFF3FA")
    set_paragraph_format(ak_hdr, keep_with_next=True)
    rk = ak_hdr.add_run("ANSWER KEY")
    rk.font.name      = DEFAULT_FONT
    rk.font.size      = Pt(FONT_HEADING)
    rk.font.bold      = True
    rk.font.color.rgb = RGB_NAVY

    add_answer_key_table(doc, answer_key)

    doc.save(str(output_path))


# ── Run all 4 ─────────────────────────────────────────────────────────────────
print()
print("=" * 60)
print("STEP 4: Generate DOCX files")
print("=" * 60)
print()

for letter, qs in exams.items():
    out = OUTPUT_DIR / f"ITE_{exam_year}_Practice_Exam_{letter}_Cole.docx"
    print(f"  Building Exam {letter} ({len(qs)} questions)...", end=" ", flush=True)
    build_exam_docx(qs, letter, out)
    print(f"✓  {out.name}")

print()
print("=" * 60)
print("COMPLETE")
print("=" * 60)
print(f"\n4 exam files written to:")
print(f"  {OUTPUT_DIR}")
print()
print("To regenerate the fixed single exam (ITE_2025_v3_Exam_Okezia_Cole.docx),")
print("run this from the scripts/ directory:")
print()
print("  node ite_report_builder_v2.js \\")
print('    "../resident_data/ITE_okezia_cole/outputs/analysis_v2_2025.json" \\')
print('    "../resident_data/ITE_okezia_cole/outputs/"')
