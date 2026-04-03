"""
build_aafp_qa.py
Builds a Q&A Word document from AAFP questions that have ITE citation overlap (File 3).
Sorted by ITE overlap count descending (highest yield first).
Uses word_doc_defaults.py for St. Luke's / ITE Intelligence styling.

Run from anywhere — paths are dynamic.
Output: 03_module.3_analyst/reports/AAFP_BRQ_ITE_Overlap_QA_v2.docx
"""

import json
import sqlite3
import sys
from pathlib import Path

# ── Paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
DB_PATH      = PROJECT_ROOT / "00_database" / "db" / "ite_intelligence.db"
DEFAULTS_DIR = SCRIPT_DIR   # word_doc_defaults.py lives in the same directory
OUT_PATH     = PROJECT_ROOT / "03_module.3_analyst" / "reports" / "AAFP_BRQ_ITE_Overlap_QA_v2.docx"

sys.path.insert(0, str(DEFAULTS_DIR))
from word_doc_defaults import (
    new_document, add_title, add_subtitle, add_section_header, add_body_text,
    add_divider, add_left_border, add_shading, add_bottom_border, add_page_number_footer,
    set_paragraph_format, sanitize,
    NAVY, GOLD, BLUE, LIGHT_BLUE, DARK_TEXT, MED_GRAY,
    RGB_NAVY, RGB_GOLD, RGB_BLUE, RGB_DARK_TEXT, RGB_GRAY, RGB_GREEN,
    DEFAULT_FONT, FONT_HEADING, FONT_BODY, FONT_SMALL, FONT_TINY,
)

from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Query ─────────────────────────────────────────────────────────────────────
QUERY = """
SELECT
    aq.aafp_qid,
    aq.stem,
    aq.choices,
    aq.correct_letter,
    aq.correct_text,
    aq.explanation,
    aq.quiz_title,
    aq.question_number,
    aq.body_system,
    aq.blueprint,
    aq.url,
    COUNT(DISTINCT qrp.qid)       AS ite_overlap_count,
    COUNT(DISTINCT ac.article_id) AS shared_article_count
FROM aafp_questions aq
JOIN aafp_citations ac
    ON ac.aafp_qid = aq.aafp_qid
JOIN articles a
    ON a.article_id = ac.article_id
JOIN question_ref_pairs qrp
    ON qrp.clean_ref = a.clean_ref
GROUP BY aq.aafp_qid
HAVING COUNT(DISTINCT qrp.qid) > 0
ORDER BY ite_overlap_count DESC, aq.aafp_qid
"""

# ── Helpers ───────────────────────────────────────────────────────────────────

def overlap_badge(count):
    if count >= 8: return f"High-Yield ITE Overlap  ({count} ITE questions)"
    if count >= 4: return f"Moderate ITE Overlap  ({count} ITE questions)"
    return f"ITE Overlap  ({count} ITE questions)"

def badge_color(count):
    if count >= 8: return RGB_GOLD
    if count >= 4: return RGB_BLUE
    return RGB_DARK_TEXT

def add_badge_line(doc, overlap_count, quiz_title, q_number, body_system):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.25)
    set_paragraph_format(p, keep_with_next=True)

    badge_run = p.add_run(overlap_badge(overlap_count))
    badge_run.font.name  = DEFAULT_FONT
    badge_run.font.size  = Pt(FONT_TINY)
    badge_run.font.bold  = True
    badge_run.font.color.rgb = badge_color(overlap_count)

    sep = p.add_run("   |   ")
    sep.font.name  = DEFAULT_FONT
    sep.font.size  = Pt(FONT_TINY)
    sep.font.color.rgb = RGB_GRAY

    info = p.add_run(f"{quiz_title}  ·  Q{q_number}  ·  {body_system or 'N/A'}")
    info.font.name  = DEFAULT_FONT
    info.font.size  = Pt(FONT_TINY)
    info.font.color.rgb = RGB_GRAY
    return p

def add_stem(doc, number, stem_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.25)
    set_paragraph_format(p, keep_with_next=True)

    num_run = p.add_run(f"Question {number}  ")
    num_run.font.name  = DEFAULT_FONT
    num_run.font.size  = Pt(FONT_HEADING)
    num_run.font.bold  = True
    num_run.font.color.rgb = RGB_NAVY

    stem_run = p.add_run(sanitize(stem_text) if stem_text else "")
    stem_run.font.name  = DEFAULT_FONT
    stem_run.font.size  = Pt(FONT_BODY)
    stem_run.font.color.rgb = RGB_DARK_TEXT
    return p

def add_choice(doc, letter, text, is_correct=False):
    """Single answer choice — is_correct unused (cover-friendly exam format)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    set_paragraph_format(p, keep_with_next=True)

    letter_run = p.add_run(f"{letter}.  ")
    letter_run.font.name  = DEFAULT_FONT
    letter_run.font.size  = Pt(FONT_BODY)
    letter_run.font.color.rgb = RGB_DARK_TEXT

    text_run = p.add_run(sanitize(text) if text else "")
    text_run.font.name  = DEFAULT_FONT
    text_run.font.size  = Pt(FONT_BODY)
    text_run.font.color.rgb = RGB_DARK_TEXT
    return p

def add_correct_answer_line(doc, letter, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.25)
    set_paragraph_format(p, keep_with_next=True)

    check_run = p.add_run("CORRECT ANSWER: ")
    check_run.font.name  = DEFAULT_FONT
    check_run.font.size  = Pt(FONT_BODY)
    check_run.font.bold  = True
    check_run.font.color.rgb = RGB_GREEN

    ans_run = p.add_run(f"{letter} — {sanitize(text) if text else ''}")
    ans_run.font.name  = DEFAULT_FONT
    ans_run.font.size  = Pt(FONT_BODY)
    ans_run.font.color.rgb = RGB_GREEN
    return p

def add_explanation(doc, text):
    if not text:
        return
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(4)
    label.paragraph_format.space_after  = Pt(2)
    label.paragraph_format.left_indent  = Inches(0.25)
    set_paragraph_format(label, keep_with_next=True)
    label_run = label.add_run("Explanation:")
    label_run.font.name  = DEFAULT_FONT
    label_run.font.size  = Pt(FONT_SMALL)
    label_run.font.bold  = True
    label_run.font.color.rgb = RGB_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    exp_run = p.add_run(sanitize(text))
    exp_run.font.name  = DEFAULT_FONT
    exp_run.font.size  = Pt(FONT_SMALL)
    exp_run.font.color.rgb = RGB_DARK_TEXT

# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    rows = conn.execute(QUERY).fetchall()
    conn.close()
    print(f"Loaded {len(rows)} AAFP questions with ITE citation overlap")

    high_yield = sum(1 for r in rows if r["ite_overlap_count"] >= 8)
    mod_yield  = sum(1 for r in rows if 4 <= r["ite_overlap_count"] < 8)

    doc = new_document()

    add_title(doc, "AAFP Board Review Questions")
    add_subtitle(doc, "ITE-Aligned Study Set — Questions with Answers & Explanations")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(16)
    stat_run = p.add_run(
        f"{len(rows)} Questions   |   "
        f"{high_yield} High-Yield (8+ ITE overlap)   |   "
        f"{mod_yield} Moderate (4–7 ITE overlap)   |   "
        f"Family Medicine Board Prep"
    )
    stat_run.font.name  = DEFAULT_FONT
    stat_run.font.size  = Pt(FONT_SMALL)
    stat_run.font.color.rgb = RGB_GRAY

    add_divider(doc)

    for i, row in enumerate(rows, start=1):
        choices = json.loads(row["choices"]) if row["choices"] else []
        correct = (row["correct_letter"] or "").strip().upper()

        if (i - 1) % 25 == 0:
            add_section_header(doc, f"Questions {i}–{min(i + 24, len(rows))}", level=1)

        add_badge_line(doc, row["ite_overlap_count"], row["quiz_title"] or "",
                       row["question_number"] or "", row["body_system"] or "")
        add_stem(doc, i, row["stem"] or "")

        for ch in choices:
            add_choice(doc, ch.get("letter", "").strip().upper(), ch.get("text", ""))

        add_correct_answer_line(doc, correct, row["correct_text"] or "")
        add_explanation(doc, row["explanation"])

        if i < len(rows):
            add_divider(doc)

    for section in doc.sections:
        add_page_number_footer(section)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Saved → {OUT_PATH}")
    print(f"  {len(rows)} questions | {high_yield} high-yield | {mod_yield} moderate")

if __name__ == "__main__":
    main()
