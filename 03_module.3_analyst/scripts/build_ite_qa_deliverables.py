#!/usr/bin/env python3
"""
build_ite_qa_deliverables.py
=============================
Generates ITE Q&A deliverable files for practice_questions/ — one DOCX and
one XLSX per exam year (2018–2025), totalling 16 files.

Output:
    01_module.1_warehouse/practice_questions/word_docs/ITE_{YEAR}_QA.docx
    01_module.1_warehouse/practice_questions/excel/ITE_{YEAR}_QA.xlsx

DOCX format:
    St. Luke's palette (word_doc_defaults.py), gold left-border question
    headers, green correct-answer highlighting on choices + banner,
    light blue answer box, explanation + reference in small/gray.
    Question header includes QID, body system, and blueprint category.

XLSX format:
    Flat table — #, QID, Year, Body System, Blueprint, stem, A–E choices
    split, correct letter + answer (green), explanation, reference.
    Frozen header, alternating row shading.

Source table: questions (ite_intelligence.db)

Usage:
    python build_ite_qa_deliverables.py              # all 16 files
    python build_ite_qa_deliverables.py --dry-run    # counts only, no files
    python build_ite_qa_deliverables.py --year 2025  # single year
    python build_ite_qa_deliverables.py --docx-only
    python build_ite_qa_deliverables.py --xlsx-only
"""

import json
import sys
import sqlite3
import argparse
from pathlib import Path

# ── Paths ──────────────────────────────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
sys.path.insert(0, str(SCRIPT_DIR))   # word_doc_defaults lives here

DB_PATH  = PROJECT_ROOT / "00_database" / "db" / "ite_intelligence.db"
OUT_DOCX = PROJECT_ROOT / "01_module.1_warehouse" / "practice_questions" / "word_docs"
OUT_XLSX = PROJECT_ROOT / "01_module.1_warehouse" / "practice_questions" / "excel"
LOG_DIR  = PROJECT_ROOT / "00_database" / "logs"

ITE_YEARS = [2018, 2019, 2020, 2021, 2022, 2023, 2024, 2025]


# ── DB helpers ─────────────────────────────────────────────────────────────────

def load_year(conn: sqlite3.Connection, year: int) -> list:
    """
    Returns list of question dicts for the given exam year, ordered by qid.
    Each dict: qid, exam_year, body_system, blueprint, question_text,
               choices (list), correct_letter, correct_text, explanation, reference
    """
    cur = conn.cursor()
    cur.execute("""
        SELECT qid, exam_year, body_system, blueprint,
               question_text, choices, correct_letter, correct_text,
               explanation, reference
        FROM questions
        WHERE exam_year = ?
        ORDER BY qid
    """, (year,))
    rows = cur.fetchall()

    questions = []
    for row in rows:
        qid, yr, body_sys, blueprint, q_text, choices_json, \
            correct_letter, correct_text, explanation, reference = row

        try:
            choices = json.loads(choices_json) if choices_json else []
        except Exception:
            choices = []

        questions.append({
            "qid":            qid or "",
            "exam_year":      yr,
            "body_system":    body_sys or "",
            "blueprint":      blueprint or "",
            "question_text":  q_text or "",
            "choices":        choices,
            "correct_letter": correct_letter or "",
            "correct_text":   correct_text or "",
            "explanation":    explanation or "",
            "reference":      reference or "",
        })
    return questions


# ── DOCX builder ───────────────────────────────────────────────────────────────

def build_docx(year: int, questions: list, out_path: Path) -> int:
    """Write one DOCX for the year. Returns question count."""
    from word_doc_defaults import (
        new_document, add_title, add_subtitle, add_divider,
        add_shading, add_left_border, add_page_number_footer,
        sanitize, set_paragraph_format,
        NAVY, GOLD, LIGHT_BLUE,
        RGB_NAVY, RGB_BLUE, RGB_DARK_TEXT, RGB_GRAY, RGB_GREEN,
        DEFAULT_FONT, FONT_HEADING, FONT_BODY, FONT_SMALL, FONT_TINY,
    )
    from docx.shared import Pt, Inches

    doc = new_document()
    add_title(doc, f"ABFM ITE — {year}")
    add_subtitle(doc, f"{len(questions)} Questions")

    # Footer: year-specific
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.clear()
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(_qn('w:val'), 'right')
    tab.set(_qn('w:pos'), '9360')
    tabs.append(tab)
    pPr.append(tabs)
    lr = p.add_run(f"{year} ABFM ITE  |  St. Luke's Family Medicine Residency")
    lr.font.name = DEFAULT_FONT; lr.font.size = Pt(FONT_TINY); lr.font.color.rgb = RGB_GRAY
    p.add_run("\t").font.size = Pt(FONT_TINY)
    pr = p.add_run("Page ")
    pr.font.name = DEFAULT_FONT; pr.font.size = Pt(FONT_TINY); pr.font.color.rgb = RGB_GRAY
    fld = p.add_run()
    fld.font.size = Pt(FONT_TINY); fld.font.color.rgb = RGB_GRAY
    fc1 = OxmlElement('w:fldChar'); fc1.set(_qn('w:fldCharType'), 'begin')
    it  = OxmlElement('w:instrText'); it.set(_qn('xml:space'), 'preserve'); it.text = 'PAGE'
    fc2 = OxmlElement('w:fldChar'); fc2.set(_qn('w:fldCharType'), 'end')
    fld._r.extend([fc1, it, fc2])

    total_q = len(questions)

    for idx, q in enumerate(questions, start=1):
        qid        = q["qid"]
        body_sys   = sanitize(q["body_system"])
        blueprint  = sanitize(q["blueprint"])
        q_text     = sanitize(q["question_text"])
        choices    = q["choices"]
        c_letter   = q["correct_letter"]
        c_text     = sanitize(q["correct_text"])
        expl       = sanitize(q["explanation"])
        ref        = sanitize(q["reference"])

        # ── Question header ────────────────────────────────────────────────────
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(14)
        p_hdr.paragraph_format.space_after  = Pt(4)
        add_left_border(p_hdr, color=GOLD, size=24)
        add_shading(p_hdr, "EFF3FA")
        set_paragraph_format(p_hdr, keep_with_next=True)

        r_num = p_hdr.add_run(f"Q{idx}  ")
        r_num.font.name = DEFAULT_FONT; r_num.font.size = Pt(FONT_HEADING)
        r_num.font.bold = True; r_num.font.color.rgb = RGB_NAVY

        meta = f"({qid}"
        if body_sys:
            meta += f"  ·  {body_sys}"
        if blueprint:
            meta += f"  ·  {blueprint}"
        meta += ")"
        r_meta = p_hdr.add_run(meta)
        r_meta.font.name = DEFAULT_FONT; r_meta.font.size = Pt(FONT_TINY)
        r_meta.font.bold = False; r_meta.font.color.rgb = RGB_GRAY

        # ── Stem ──────────────────────────────────────────────────────────────
        p_stem = doc.add_paragraph()
        p_stem.paragraph_format.left_indent  = Inches(0.15)
        p_stem.paragraph_format.space_before = Pt(4)
        p_stem.paragraph_format.space_after  = Pt(6)
        set_paragraph_format(p_stem, keep_with_next=True)
        r_stem = p_stem.add_run(q_text)
        r_stem.font.name = DEFAULT_FONT; r_stem.font.size = Pt(FONT_BODY)
        r_stem.font.color.rgb = RGB_DARK_TEXT

        # ── Answer choices ─────────────────────────────────────────────────────
        for ch in choices:
            letter     = ch.get("letter", "")
            ch_text    = sanitize(ch.get("text", ""))
            is_correct = (letter == c_letter)

            p_ch = doc.add_paragraph()
            p_ch.paragraph_format.left_indent  = Inches(0.4)
            p_ch.paragraph_format.space_before = Pt(1)
            p_ch.paragraph_format.space_after  = Pt(1)
            set_paragraph_format(p_ch, keep_with_next=True)

            r_ltr = p_ch.add_run(f"{letter})  ")
            r_ltr.font.name = DEFAULT_FONT; r_ltr.font.size = Pt(FONT_BODY)
            r_ltr.font.bold = False
            r_ltr.font.color.rgb = RGB_DARK_TEXT

            r_cht = p_ch.add_run(ch_text)
            r_cht.font.name = DEFAULT_FONT; r_cht.font.size = Pt(FONT_BODY)
            r_cht.font.bold = False
            r_cht.font.color.rgb = RGB_DARK_TEXT

        # ── Correct answer banner ──────────────────────────────────────────────
        p_ans = doc.add_paragraph()
        p_ans.paragraph_format.left_indent  = Inches(0.15)
        p_ans.paragraph_format.space_before = Pt(8)
        p_ans.paragraph_format.space_after  = Pt(2)
        add_shading(p_ans, LIGHT_BLUE)
        set_paragraph_format(p_ans, keep_with_next=True)

        r_lbl = p_ans.add_run("✓ Answer: ")
        r_lbl.font.name = DEFAULT_FONT; r_lbl.font.size = Pt(FONT_BODY)
        r_lbl.font.bold = True; r_lbl.font.color.rgb = RGB_GREEN

        r_aval = p_ans.add_run(f"{c_letter})  {c_text}")
        r_aval.font.name = DEFAULT_FONT; r_aval.font.size = Pt(FONT_BODY)
        r_aval.font.color.rgb = RGB_DARK_TEXT

        # ── Explanation ────────────────────────────────────────────────────────
        if expl:
            p_expl = doc.add_paragraph()
            p_expl.paragraph_format.left_indent  = Inches(0.15)
            p_expl.paragraph_format.space_before = Pt(4)
            p_expl.paragraph_format.space_after  = Pt(4)
            set_paragraph_format(p_expl)
            r_el = p_expl.add_run("Explanation:  ")
            r_el.font.name = DEFAULT_FONT; r_el.font.size = Pt(FONT_SMALL)
            r_el.font.bold = True; r_el.font.color.rgb = RGB_BLUE
            r_et = p_expl.add_run(expl)
            r_et.font.name = DEFAULT_FONT; r_et.font.size = Pt(FONT_SMALL)
            r_et.font.color.rgb = RGB_DARK_TEXT

        # ── Reference ─────────────────────────────────────────────────────────
        if ref:
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.left_indent  = Inches(0.15)
            p_ref.paragraph_format.space_before = Pt(2)
            p_ref.paragraph_format.space_after  = Pt(8)
            set_paragraph_format(p_ref)
            r_rl = p_ref.add_run("Ref:  ")
            r_rl.font.name = DEFAULT_FONT; r_rl.font.size = Pt(FONT_TINY)
            r_rl.font.bold = True; r_rl.font.color.rgb = RGB_GRAY
            r_rt = p_ref.add_run(ref)
            r_rt.font.name = DEFAULT_FONT; r_rt.font.size = Pt(FONT_TINY)
            r_rt.font.italic = True; r_rt.font.color.rgb = RGB_GRAY

        # ── Divider (not after last question) ──────────────────────────────────
        if idx < total_q:
            add_divider(doc)

    doc.save(str(out_path))
    return total_q


# ── XLSX builder ───────────────────────────────────────────────────────────────

def build_xlsx(year: int, questions: list, out_path: Path) -> int:
    """Write one XLSX for the year. Returns question count."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    from word_doc_defaults import sanitize

    HEADER_FILL  = PatternFill("solid", fgColor="1B3564")  # navy
    CORRECT_FILL = PatternFill("solid", fgColor="E6F4EA")  # light green
    ALT_FILL     = PatternFill("solid", fgColor="F5F8FC")  # very light blue

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"ITE {year}"

    headers = ["#", "QID", "Year", "Stem", "A", "B", "C", "D", "E",
               "Correct", "Correct Answer", "Explanation", "Reference"]
    ws.append(headers)

    # Header row styling
    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font      = Font(name="Aptos", bold=True, color="FFFFFF", size=10)
        cell.fill      = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # Column widths: #, QID, Year, Stem, A-E, Correct, CorrectAns, Expl, Ref
    col_widths = [5, 16, 6, 55, 22, 22, 22, 22, 22, 8, 28, 60, 55]
    for i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws.row_dimensions[1].height = 22

    for row_idx, q in enumerate(questions, start=1):
        choices = {ch["letter"]: sanitize(ch.get("text", "")) for ch in q["choices"]}
        c_ltr   = q["correct_letter"]

        data_row = [
            row_idx,
            q["qid"],
            q["exam_year"],
            sanitize(q["question_text"]),
            choices.get("A", ""),
            choices.get("B", ""),
            choices.get("C", ""),
            choices.get("D", ""),
            choices.get("E", ""),
            c_ltr,
            sanitize(q["correct_text"]),
            sanitize(q["explanation"]),
            sanitize(q["reference"]),
        ]
        ws.append(data_row)

        excel_row = row_idx + 1
        alt = (row_idx % 2 == 0)

        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=excel_row, column=col_idx)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if col_idx in (10, 11):   # Correct / Correct Answer
                cell.font = Font(name="Aptos", size=9, bold=True, color="3E8D27")
                cell.fill = CORRECT_FILL
            else:
                cell.font = Font(name="Aptos", size=9)
                if alt:
                    cell.fill = ALT_FILL

    ws.freeze_panes = "A2"
    wb.save(str(out_path))
    return len(questions)


# ── Main ────────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(
        description="Generate ITE Q&A deliverable files — 1 DOCX + 1 XLSX per exam year"
    )
    ap.add_argument("--dry-run",   action="store_true",
                    help="Print counts only — no files written")
    ap.add_argument("--year",      type=int, default=None,
                    help="Generate for a single year only (e.g. --year 2025)")
    ap.add_argument("--docx-only", action="store_true",
                    help="Write DOCX files only (skip XLSX)")
    ap.add_argument("--xlsx-only", action="store_true",
                    help="Write XLSX files only (skip DOCX)")
    args = ap.parse_args()

    do_docx = not args.xlsx_only
    do_xlsx = not args.docx_only

    target_years = [args.year] if args.year else ITE_YEARS

    print(f"\n{'='*60}")
    print(f"  build_ite_qa_deliverables.py")
    print(f"  years={target_years}  dry-run={args.dry_run}")
    print(f"  docx={do_docx}  xlsx={do_xlsx}")
    print(f"{'='*60}\n")

    conn = sqlite3.connect(DB_PATH)

    total_files = 0
    year_counts = {}
    for year in target_years:
        questions = load_year(conn, year)
        year_counts[year] = len(questions)
        print(f"  {year}: {len(questions):3d} questions  →  "
              f"ITE_{year}_QA.docx  +  ITE_{year}_QA.xlsx")

    conn.close()

    print(f"\n  Total: {sum(year_counts.values())} questions across "
          f"{len(target_years)} year(s)")
    print(f"  Files to generate: {len(target_years) * (2 if do_docx and do_xlsx else 1)}")

    if args.dry_run:
        print(f"\n  [DRY RUN] No files written.")
        print(f"{'='*60}\n")
        return

    # ── Output dirs ─────────────────────────────────────────────────────────────
    if do_docx:
        OUT_DOCX.mkdir(parents=True, exist_ok=True)
    if do_xlsx:
        OUT_XLSX.mkdir(parents=True, exist_ok=True)

    # ── Generate ─────────────────────────────────────────────────────────────────
    conn = sqlite3.connect(DB_PATH)
    for year in target_years:
        questions = load_year(conn, year)
        stem = f"ITE_{year}_QA"

        if do_docx:
            out = OUT_DOCX / f"{stem}.docx"
            n = build_docx(year, questions, out)
            total_files += 1
            print(f"  ✓ {stem}.docx  ({n} Q)")

        if do_xlsx:
            out = OUT_XLSX / f"{stem}.xlsx"
            n = build_xlsx(year, questions, out)
            total_files += 1
            print(f"  ✓ {stem}.xlsx  ({n} Q)")

    conn.close()

    print(f"\n  Done — {total_files} files written.")
    print(f"  DOCX → {OUT_DOCX}")
    print(f"  XLSX → {OUT_XLSX}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
