#!/usr/bin/env python3
"""
build_resident_guide.py — Resident Interpretation Guide (DOCX)
===============================================================
Generates a comprehensive DOCX guide for residents to interpret ITE Score Analysis reports.
Uses word_doc_defaults.py for styling and layout.

Usage:
    python3 build_resident_guide.py

Output:
    ../docs/ITE_Report_Guide_Resident_v2.docx
"""

from pathlib import Path
import sys
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

from word_doc_defaults import *

OUTPUT_DIR = SCRIPT_DIR.parent / "docs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "ITE_Report_Guide_Resident_v2.docx"

# ==============================================================================
# CUSTOM HELPERS
# ==============================================================================

def add_action_callout(doc, text):
    """Gold-bordered ACTION callout with light gold background."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_left_border(p, color=GOLD, size=24)
    add_shading(p, "FEF9EC")

    label = p.add_run("✎ ACTION  ")
    label.font.name = DEFAULT_FONT
    label.font.size = Pt(FONT_BODY)
    label.font.bold = True
    label.font.color.rgb = RGB_GOLD

    body = p.add_run(text)
    body.font.name = DEFAULT_FONT
    body.font.size = Pt(FONT_BODY)
    body.font.color.rgb = RGB_DARK_TEXT

    return p

def add_bullet(doc, text, indent=0.5):
    """Formatted bullet list item."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(FONT_BODY)
    run.font.color.rgb = RGB_DARK_TEXT
    return p

def add_numbered(doc, text, indent=0.5):
    """Formatted numbered list item."""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(FONT_BODY)
    run.font.color.rgb = RGB_DARK_TEXT
    return p

# ==============================================================================
# BUILD DOCUMENT
# ==============================================================================

def build_guide():
    """Construct the resident interpretation guide."""
    doc = new_document()

    # ========== COVER PAGE ==========
    add_title(doc, "ITE Score Analysis")
    add_subtitle(doc, "Resident's Reading Guide")
    add_subtitle(doc, "How to turn your report into a study plan")
    add_divider(doc)
    add_body_text(
        doc,
        "The ITE Score Analysis report is designed to show you exactly where your "
        "knowledge gaps are and which gaps are worth your study time. This guide walks "
        "through the report section by section. For each section, you will learn what you "
        "are reading and what to do with it."
    )

    doc.add_page_break()

    # ========== SECTION 0: EXAM AT A GLANCE ==========
    add_section_header(doc, "Section 0 — Exam at a Glance", level=1)
    add_body_text(
        doc,
        "National statistics about the exam itself — how many items are scored, the passing "
        "score, and where the average resident in your PGY year scored."
    )

    add_section_header(doc, "Key things to note", level=2)
    add_bullet(doc, "MPS is 380. Below 380 = your #1 goal is closing this gap before the FMCE.")
    add_bullet(doc, "440 is the FMCE signal threshold. If you are comfortably above MPS but below 440, that is your target zone.")
    add_bullet(doc, "SEM +/-38 means there is natural measurement error. A score of 375 and 380 are not meaningfully different — they are within one measurement unit of each other.")
    add_bullet(doc, "The experimental items (deleted items) do not count toward your score. Do not be confused by question numbers that do not appear in your score.")

    add_action_callout(
        doc,
        "Before reading further, note your score against these three marks: MPS (380), "
        "FMCE threshold (440), and the listed PGY Mean for your year. These three numbers "
        "anchor everything that follows."
    )

    # ========== SECTION 1: YOUR SCORE ==========
    add_section_header(doc, "Section 1 — Your Score", level=1)
    add_body_text(
        doc,
        "Three versions of your score — raw %, scaled score, and percentile — plus comparison "
        "to the national average for your PGY level."
    )

    add_section_header(doc, "The numbers explained", level=2)
    add_bullet(doc, "Raw %: The fraction of questions you answered correctly out of the total scored items.")
    add_bullet(doc, "Scaled score: A standardized number (roughly 200-600) that accounts for year-to-year variation in difficulty. This is the number that matters for passing (MPS = 380) and the FMCE (440 signal).")
    add_bullet(doc, "Percentile: Where you rank among all residents who took that year's exam.")

    add_body_text(
        doc,
        "The confidence range (68%) shows that your actual knowledge level falls somewhere within "
        "this band. A scored test always has measurement error — this range captures it."
    )

    add_body_text(
        doc,
        "If you see a trajectory label like STRONG or AT RISK, treat it as a directional signal, "
        "not a verdict. It is an estimate based on your score relative to national benchmarks."
    )

    add_action_callout(
        doc,
        "Look at vs MPS and vs PGY Mean. These two numbers tell you the whole story. If vs MPS is "
        "negative, getting to 380 is your priority. If you are above MPS, pushing toward 440 is the goal."
    )

    # ========== SECTION 2: EXAM SUMMARY ==========
    add_section_header(doc, "Section 2 — Exam Summary", level=1)
    add_body_text(
        doc,
        "A quick-reference bullet list summarizing your most important findings."
    )

    add_section_header(doc, "The bullets explained", level=2)
    add_bullet(doc, "Scaled score line: Same as Section 1 — confirms your baseline.")
    add_bullet(doc, "X additional correct answers needed to reach MPS: A rough estimate of how many more questions you would need to answer correctly to hit 380. Use it for motivation, not precision.")
    add_bullet(doc, "Weak areas: The content domains where your score is below your personal average OR below 70%. These are your study priorities for Part B and the reading list.")
    add_bullet(doc, "Lowest blueprint category: The single domain where you struggled most. Start your content review here.")
    add_bullet(doc, "Easy misses: The most important number in this section. These are questions most of your peers got right that you missed. See Section 6 for the full breakdown.")

    add_action_callout(
        doc,
        "Highlight the Lowest blueprint category and Easy misses count. Those two facts should "
        "anchor your study plan."
    )

    # ========== SECTION 3B: YEAR-OVER-YEAR PROGRESS ==========
    add_section_header(doc, "Section 3b — Year-Over-Year Progress", level=1)
    add_body_text(
        doc,
        "(Only appears if you have a prior year's report) A direct comparison of your score and "
        "weak areas versus last year and the year before."
    )

    add_section_header(doc, "The trajectory categories", level=2)
    add_bullet(doc, "Checkmark Closed: A gap you had last year that you have resolved. Do not abandon these topics entirely — they can regress if you stop reviewing.")
    add_bullet(doc, "Warning Persistent: A gap that appeared last year and is still here. This is a signal that studying the same way is not working for this topic. You need a different approach — not just more hours, but a different strategy.")
    add_bullet(doc, "Arrow New: A gap that was not here last year. This may be new exam content, a topic you de-prioritized, or a domain where your knowledge decayed.")

    add_action_callout(
        doc,
        "Focus your energy on Persistent gaps first. If you have reviewed the same material twice "
        "and it is not sticking, change the resource or method. New gaps are secondary — they are "
        "easier to fix."
    )

    doc.add_page_break()

    # ========== SECTION 4: BLUEPRINT & BODY SYSTEM ==========
    add_section_header(doc, "Section 4 — Blueprint & Body System Performance", level=1)
    add_body_text(
        doc,
        "Your performance broken down by ABFM content domain (Blueprint) and by organ system "
        "(Body System)."
    )

    add_section_header(doc, "How to read the Blueprint table", level=2)
    add_bullet(doc, "Prior %: Your rate last year (dashes = first year or no prior data).")
    add_bullet(doc, "Delta: Change from last year. Green = improved, Red = regressed.")
    add_bullet(doc, "SEM: The measurement error for this sub-score. A large SEM (e.g., +/-15%) means this sub-score is based on few questions and is less reliable. Treat large-SEM sub-scores as hypotheses to investigate, not confirmed deficits.")
    add_bullet(doc, "Color coding: Green (70%+) = passing territory. Amber (50-70%) = needs work. Red (<50%) = urgent.")

    add_body_text(
        doc,
        "The two Body System tables: ABFM-Reported (navy label) came directly from your official "
        "ABFM score report — these are the most reliable sub-scores. Database-Derived (blue label) "
        "is inferred from the content of your exam questions using the ITE Knowledge Base. These are "
        "supplementary signals — useful for pattern detection, not for precise scoring."
    )

    add_action_callout(
        doc,
        "Any row below 50% (red) is urgent. Any row in amber (50-70%) needs focused review. Within "
        "the red rows, prioritize those with a smaller SEM value — those are more reliable signals of "
        "a genuine gap. Start with the highest-priority red row and do not move to the next until you "
        "have done a content review of that domain."
    )

    # ========== SECTION 6: DIFFICULTY PROFILE ==========
    add_section_header(doc, "Section 6 — Difficulty Profile", level=1)
    add_body_text(
        doc,
        "Your missed questions broken down by how difficult they were nationally."
    )

    add_section_header(doc, "The three tiers", level=2)
    add_bullet(doc, "Easy Miss (700+): Nearly everyone else got these right. These are NOT unlucky hard questions — they are genuine knowledge gaps. This is where your study time will produce the most improvement.")
    add_bullet(doc, "Mid-Range (300-699): Challenging questions that reward high-yield review. Each one you convert is meaningful.")
    add_bullet(doc, "Hard Miss (<300): Most residents missed these too. Chasing them is low ROI unless you have already cleaned up the Easy and Mid-Range tiers.")

    add_body_text(
        doc,
        "The Easy Misses table lists each specific question (QID) that was an Easy Miss. These are "
        "your highest-priority review items. Look each one up in your exam materials and understand "
        "the concept it tested."
    )

    add_action_callout(
        doc,
        "Print or note the Easy Misses table. Each QID listed is a question where you and most of "
        "your peers diverged — meaning it is a learnable, reviewable gap. Work through every one "
        "before your next exam."
    )

    # ========== SECTION 7: LOW-HANGING FRUIT ==========
    add_section_header(doc, "Section 7 — Low-Hanging Fruit", level=1)
    add_body_text(
        doc,
        "A ranked list of content domains by study ROI — where each hour of review converts to the "
        "most scaled score improvement."
    )

    add_section_header(doc, "The two metrics explained", level=2)
    add_bullet(doc, "Improvable Items: The number of questions in this category that (1) you got wrong AND (2) most of your peers got right. These are the questions you could realistically convert on a re-test with targeted review.")
    add_bullet(doc, "Priority Score: Improvable Items x how much this domain counts on the exam. A domain worth 35% of the exam with 3 improvable items ranks higher than a domain worth 5% with 5 improvable items.")

    add_body_text(
        doc,
        "The Blueprint table includes a Priority Score (tied to official exam weighting). The Body "
        "System table shows only Improvable Items (body systems do not have published ABFM weightings)."
    )

    add_action_callout(
        doc,
        "Start with Rank #1 on the Blueprint table. Spend concentrated study time there before moving "
        "to Rank #2. This section is your study sequence generator — follow it."
    )

    # ========== SECTION 7B: CATEGORY CROSSOVER ==========
    add_section_header(doc, "Section 7b — Category Crossover Weaknesses", level=1)
    add_body_text(
        doc,
        "Weak spots at the intersection of TWO dimensions — for example, Cardiovascular x Acute Care "
        "means acute presentations of cardiovascular disease specifically (not cardiovascular disease "
        "broadly, not acute care broadly — the overlap)."
    )

    add_section_header(doc, "Why this matters", level=2)
    add_body_text(
        doc,
        "A pure Cardiovascular weakness might mean you need to review all cardiac medicine. A "
        "Cardiovascular x Acute Care crossover means you specifically need to focus on acute cardiac "
        "presentations (chest pain, ACS, arrhythmia emergencies) — a much tighter study target."
    )

    add_action_callout(
        doc,
        "If any crossover weakness aligns with one of your red/amber Blueprint rows from Section 4, "
        "that intersection is your most specific study target. Narrow your review to clinical scenarios "
        "at exactly that overlap."
    )

    doc.add_page_break()

    # ========== SECTION 8: CONCEPT FINGERPRINT ==========
    add_section_header(doc, "Section 8 — Concept Fingerprint", level=1)
    add_body_text(
        doc,
        "Drugs that appeared repeatedly across your missed questions — a fingerprint of recurring "
        "pharmacology knowledge gaps."
    )

    add_section_header(doc, "How to read the frequency column", level=2)
    add_bullet(doc, "3x+ (red): This drug appeared in 3 or more of your missed questions. This is a consistent pharmacology gap — not a coincidence.")
    add_bullet(doc, "2x (amber): Two missed questions shared this drug. Worth targeted review.")
    add_bullet(doc, "The QID column: Lists the specific question IDs where this drug appeared. Use them to pull those questions from your exam materials.")

    add_body_text(
        doc,
        "Badges (appear only on multi-year reports): A Persistent badge means this drug cluster was "
        "in your missed questions last year too. A New badge means it is new this year."
    )

    add_action_callout(
        doc,
        "For any drug appearing 2+ times, review: mechanism of action, indications, contraindications, "
        "key drug interactions, and monitoring parameters. Use the QID column to find those specific "
        "questions and understand why you missed them."
    )

    # ========== SECTION 9: ICD-10 WEAKNESS MAP ==========
    add_section_header(doc, "Section 9 — ICD-10 Weakness Map", level=1)
    add_body_text(
        doc,
        "Your missed questions mapped to clinical conditions by ICD-10 code, showing which diseases "
        "and syndromes appear most frequently in your knowledge gaps."
    )

    add_section_header(doc, "How to read it", level=2)
    add_body_text(
        doc,
        "The Misses column shows how many of your missed questions involve each condition. Red (3+) "
        "means this condition appeared across multiple exam questions that you missed — it is a "
        "recurring exam theme."
    )

    add_body_text(
        doc,
        "The Clinical Domain column groups codes by organ system (e.g., Circulatory, "
        "Endocrine/Metabolic) — useful for seeing whether gaps cluster in a single specialty area."
    )

    add_action_callout(
        doc,
        "For any condition with 3+ misses, go to the matching article in your High-Yield Reading "
        "List (Section 10). If no article is listed for it, search your clinical guidelines or "
        "UpToDate for that specific condition."
    )

    # ========== SECTION 10: HIGH-YIELD READING LIST ==========
    add_section_header(doc, "Section 10 — High-Yield Reading List", level=1)
    add_body_text(
        doc,
        "A curated reading list in two tiers, drawn from the ITE Knowledge Base of 1,998 clinical "
        "guidelines and review articles."
    )

    add_section_header(doc, "The two tiers", level=2)
    add_bullet(doc, "Targeted to Your Exam (navy header): Articles specifically linked to questions you missed on this exam. These articles cover the clinical material behind your actual exam misses. Read these first.")
    add_bullet(doc, "High-Yield for All Residents (blue header): The most frequently cited articles in the ABFM exam database — cornerstone references that appear repeatedly across all FM residents exams. These are foundational reading regardless of your specific gaps.")

    add_section_header(doc, "The summary table columns", level=2)
    add_bullet(doc, "Citations: How many times this article has been cited across all ITE years. Higher = more likely to appear on future exams.")
    add_bullet(doc, "Exam Yrs: How many different exam years referenced this article. Cited in 4+ years = very stable exam topic.")
    add_bullet(doc, "Weak Links: How many of your specific missed questions this article covers.")
    add_bullet(doc, "Status: Whether the guideline is current. Current = safe to use. Updated = a newer guideline may exist — verify before relying on it.")

    add_body_text(
        doc,
        "The QID glossary under each article shows the exact questions from your exam that this "
        "article covers. If you read that article, you are directly addressing those question topics."
    )

    add_action_callout(
        doc,
        "Read every article in the Targeted tier before your next exam. Use the QID glossary to "
        "connect each article back to your actual exam misses. For High-Yield articles, start with "
        "the highest citation count and work down."
    )

    doc.add_page_break()

    # ========== PART B: PRACTICE QUESTIONS ==========
    add_section_header(doc, "Part B — Practice Questions", level=1)
    add_body_text(
        doc,
        "A curated set of practice questions from the ITE + AAFP BRQ database, selected specifically "
        "for your weak areas."
    )

    add_section_header(doc, "How they were selected", level=2)
    add_body_text(
        doc,
        "Questions were chosen by matching your weak areas (from Sections 4, 6, 7) to the ITE "
        "Knowledge Base. The Targeting column shows which weak area drove the selection:"
    )
    add_bullet(doc, "Blue text = Blueprint category match (e.g., Chronic Care)")
    add_bullet(doc, "Green text = Body system match (e.g., Cardiovascular)")
    add_bullet(doc, "Purple text = Crossover match (e.g., Cardiovascular x Acute Care)")
    add_bullet(doc, "Amber text = Concept fingerprint match")

    add_body_text(
        doc,
        "The Match column shows how tightly matched the question is: Direct Match = the question is "
        "closely linked to one of your actual missed topics. ICD-10 Sibling = the question tests a "
        "clinically related condition within the same disease family. Vector Match = the question is "
        "semantically similar to your weak areas."
    )

    add_body_text(
        doc,
        "Questions are ordered by relevance — #1 is the highest priority for your specific gaps."
    )

    add_action_callout(
        doc,
        "Complete all practice questions BEFORE checking answers. Treat it like a mini-exam. Then "
        "review the explanation for every question you got wrong — not just to get the right answer, "
        "but to identify the underlying knowledge gap. The goal is to build a pattern, not memorize "
        "individual answers."
    )

    # ========== APPENDIX: MISSED ITEMS ==========
    add_section_header(doc, "Appendix — Missed Exam Items", level=1)
    add_body_text(
        doc,
        "A reference list of every question you missed on the actual exam, with QID and blueprint/body "
        "system labels. This lets you look up any specific question in your exam materials."
    )

    add_action_callout(
        doc,
        "Use the QIDs here to cross-reference with your actual exam document (provided separately). "
        "If a missed item was also listed as an Easy Miss in Section 6, prioritize it for review."
    )

    doc.add_page_break()

    # ========== BUILDING YOUR STUDY PLAN ==========
    add_section_header(doc, "Building Your Study Plan", level=1)
    add_body_text(
        doc,
        "Synthesize the report into a prioritized action sequence:"
    )

    add_numbered(doc, "Address Easy Misses (Section 6 — QIDs listed)")
    add_numbered(doc, "Work Rank #1 from Low-Hanging Fruit (Section 7)")
    add_numbered(doc, "Read all Targeted articles (Section 10)")
    add_numbered(doc, "Complete practice questions in order (Part B)")
    add_numbered(doc, "Re-examine Persistent gaps (Section 3b) with a different study approach")
    add_numbered(doc, "Read High-Yield articles in citation-count order (Section 10)")
    add_numbered(doc, "Review concept fingerprint drugs (Section 8)")

    add_divider(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(
        "Remember: This guide is a companion to your ITE Score Analysis report. Use them together "
        "to build a focused study plan that addresses your specific knowledge gaps."
    )
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(FONT_BODY)
    run.font.italic = True
    run.font.color.rgb = RGB_DARK_TEXT

    # ========== FOOTER ==========
    add_page_number_footer(doc.sections[0])

    return doc

# ==============================================================================
# MAIN
# ==============================================================================

if __name__ == "__main__":
    doc = build_guide()
    doc.save(str(OUTPUT_PATH))
    print("✓ Resident Guide (Python) generated successfully")
    print(f"  Output: {OUTPUT_PATH}")
