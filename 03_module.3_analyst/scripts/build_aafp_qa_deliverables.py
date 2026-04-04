#!/usr/bin/env python3
"""
build_aafp_qa_deliverables.py
==============================
Generates AAFP Board Review Q&A deliverable files for practice_questions/.

Produces 13 DOCX + 13 XLSX files, chunked by quiz (~100 Q per file, whole
quizzes only — never split a quiz across files).

Output:
    01_module.1_warehouse/practice_questions/word_docs/AAFP_quiz_NNN-NNN.docx
    01_module.1_warehouse/practice_questions/excel/AAFP_quiz_NNN-NNN.xlsx

Naming:
    Display number extracted from quiz_title ("Board Review Questions NN").
    Series 1 (assessment_id <= 13980): title number as-is (001–099)
    Series 2 (assessment_id >= 13981): title number + 99 (100–136)
    File named by display range of first and last quiz in chunk.

Chunking:
    Greedy: accumulate quizzes until >= 100 questions, then close chunk.
    Produces 13 files: 103/100/100/100/100/100/100/100/109/100/100/100/9 Q.

Source tables:
    aafp_questions     — stem, choices (JSON), correct_letter, correct_text, explanation
    aafp_citation_raw  — raw_text (all citations per question joined for display)

Usage:
    python build_aafp_qa_deliverables.py
    python build_aafp_qa_deliverables.py --dry-run
    python build_aafp_qa_deliverables.py --docx-only
    python build_aafp_qa_deliverables.py --xlsx-only
"""

import re
import sys
import json
import sqlite3
import argparse
from pathlib import Path
from collections import OrderedDict
from datetime import datetime, timezone

# ── Paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
sys.path.insert(0, str(SCRIPT_DIR))   # word_doc_defaults lives here

DB_PATH  = PROJECT_ROOT / "00_database" / "db" / "ite_intelligence.db"
OUT_DOCX = PROJECT_ROOT / "01_module.1_warehouse" / "practice_questions" / "word_docs"
OUT_XLSX = PROJECT_ROOT / "01_module.1_warehouse" / "practice_questions" / "excel"
LOG_DIR  = PROJECT_ROOT / "00_database" / "logs"

S2_THRESHOLD      = 13981   # assessment_id where Series 2 begins
S2_OFFSET         = 99      # add to quiz title number for Series 2 display
TARGET_Q_PER_FILE = 100


# ── DB helpers ────────────────────────────────────────────────────────────────

def _quiz_display_num(assessment_id: int, quiz_title: str) -> int:
    """Extract numeric display number from quiz_title, apply S2 offset if needed."""
    m = re.search(r'(\d+)\s*$', quiz_title)
    base = int(m.group(1)) if m else 0
    return base + S2_OFFSET if assessment_id >= S2_THRESHOLD else base


def load_all_data(conn: sqlite3.Connection) -> tuple:
    """
    Returns:
        quizzes  — list of (assessment_id, info_dict) ordered by assessment_id
                   info_dict = {quiz_title, display_num, qids: [aafp_qid, ...]}
        q_data   — {aafp_qid: {assessment_id, quiz_title, question_number,
                                stem, choices, correct_letter, correct_text,
                                explanation, citations}}
    """
    cur = conn.cursor()

    cur.execute("""
        SELECT aafp_qid, assessment_id, quiz_title, question_number,
               stem, choices, correct_letter, correct_text, explanation
        FROM aafp_questions
        ORDER BY assessment_id, question_number
    """)
    all_rows = cur.fetchall()

    cur.execute("SELECT aafp_qid, raw_text FROM aafp_citation_raw ORDER BY citation_id")
    citations = {}
    for aafp_qid, raw_text in cur.fetchall():
        citations.setdefault(aafp_qid, []).append(raw_text or "")

    quiz_map = OrderedDict()
    q_data   = {}

    for row in all_rows:
        aafp_qid, assessment_id, quiz_title, q_num, \
            stem, choices_json, correct_letter, correct_text, explanation = row

        if assessment_id not in quiz_map:
            quiz_map[assessment_id] = {
                "quiz_title":  quiz_title,
                "display_num": _quiz_display_num(assessment_id, quiz_title),
                "qids":        [],
            }
        quiz_map[assessment_id]["qids"].append(aafp_qid)

        try:
            choices = json.loads(choices_json) if choices_json else []
        except Exception:
            choices = []

        q_data[aafp_qid] = {
            "assessment_id":   assessment_id,
            "quiz_title":      quiz_title,
            "question_number": q_num,
            "stem":            stem or "",
            "choices":         choices,
            "correct_letter":  correct_letter or "",
            "correct_text":    correct_text or "",
            "explanation":     explanation or "",
            "citations":       citations.get(aafp_qid, []),
        }

    return list(quiz_map.items()), q_data


# ── Chunking ──────────────────────────────────────────────────────────────────

def chunk_quizzes(quizzes: list) -> list:
    """
    Greedy: accumulate quizzes until cumulative Q count >= TARGET_Q_PER_FILE,
    then close the chunk. Never splits a quiz across files.
    Returns list of chunks; each chunk = [(assessment_id, info_dict), ...].
    """
    chunks        = []
    current_chunk = []
    current_count = 0

    for assessment_id, info in quizzes:
        current_chunk.append((assessment_id, info))
        current_count += len(info["qids"])
        if current_count >= TARGET_Q_PER_FILE:
            chunks.append(current_chunk)
            current_chunk = []
            current_count = 0

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def chunk_filename(chunk: list) -> str:
    """Returns stem like 'AAFP_quiz_001-023' from a chunk."""
    first = chunk[0][1]["display_num"]
    last  = chunk[-1][1]["display_num"]
    if first == last:
        return f"AAFP_quiz_{first:03d}"
    return f"AAFP_quiz_{first:03d}-{last:03d}"


def chunk_q_count(chunk: list) -> int:
    return sum(len(info["qids"]) for _, info in chunk)


# ── DOCX builder ──────────────────────────────────────────────────────────────

def build_docx(chunk: list, q_data: dict, out_path: Path) -> int:
    """Write one DOCX for the chunk. Returns question count."""
    from word_doc_defaults import (
        new_document, add_title, add_subtitle, add_section_header,
        add_body_text, add_divider, add_shading, add_left_border,
        add_page_number_footer, sanitize, set_paragraph_format,
        NAVY, GOLD, BLUE, LIGHT_BLUE,
        RGB_NAVY, RGB_BLUE, RGB_DARK_TEXT, RGB_GRAY, RGB_GREEN,
        DEFAULT_FONT, FONT_HEADING, FONT_BODY, FONT_SMALL, FONT_TINY,
    )
    from docx.shared import Pt, Inches

    first_num = chunk[0][1]["display_num"]
    last_num  = chunk[-1][1]["display_num"]
    title_str    = "AAFP Board Review Questions"
    subtitle_str = (f"Quiz {first_num:03d}–{last_num:03d}"
                    if first_num != last_num else f"Quiz {first_num:03d}")

    doc = new_document()
    add_title(doc, title_str)
    add_subtitle(doc, subtitle_str)
    add_page_number_footer(doc.sections[0])

    # Ordered list of all QIDs in this chunk
    qids_ordered = []
    for _, info in chunk:
        qids_ordered.extend(info["qids"])

    total_q = len(qids_ordered)

    for idx, aafp_qid in enumerate(qids_ordered, start=1):
        q = q_data[aafp_qid]

        stem     = sanitize(q["stem"])
        choices  = q["choices"]
        c_letter = q["correct_letter"]
        c_text   = sanitize(q["correct_text"])
        expl     = sanitize(q["explanation"])
        cits     = [sanitize(c) for c in q["citations"]]

        # ── Question header ──────────────────────────────────────────────────
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(14)
        p_hdr.paragraph_format.space_after  = Pt(4)
        add_left_border(p_hdr, color=GOLD, size=24)
        add_shading(p_hdr, "EFF3FA")
        set_paragraph_format(p_hdr, keep_with_next=True)

        r_num = p_hdr.add_run(f"Q{idx}  ")
        r_num.font.name  = DEFAULT_FONT
        r_num.font.size  = Pt(FONT_HEADING)
        r_num.font.bold  = True
        r_num.font.color.rgb = RGB_NAVY

        r_quiz = p_hdr.add_run(f"({q['quiz_title']})")
        r_quiz.font.name  = DEFAULT_FONT
        r_quiz.font.size  = Pt(FONT_SMALL)
        r_quiz.font.bold  = False
        r_quiz.font.color.rgb = RGB_GRAY

        # ── Stem ─────────────────────────────────────────────────────────────
        p_stem = doc.add_paragraph()
        p_stem.paragraph_format.left_indent  = Inches(0.15)
        p_stem.paragraph_format.space_before = Pt(4)
        p_stem.paragraph_format.space_after  = Pt(6)
        set_paragraph_format(p_stem, keep_with_next=True)
        r_stem = p_stem.add_run(stem)
        r_stem.font.name  = DEFAULT_FONT
        r_stem.font.size  = Pt(FONT_BODY)
        r_stem.font.color.rgb = RGB_DARK_TEXT

        # ── Answer choices ───────────────────────────────────────────────────
        for ch in choices:
            letter     = ch.get("letter", "")
            ch_text    = sanitize(ch.get("text", ""))
            is_correct = (letter == c_letter)

            p_ch = doc.add_paragraph()
            p_ch.paragraph_format.left_indent  = Inches(0.4)
            p_ch.paragraph_format.space_before = Pt(1)
            p_ch.paragraph_format.space_after  = Pt(1)
            set_paragraph_format(p_ch, keep_with_next=True)

            r_ltr = p_ch.add_run(f"{letter})  ")
            r_ltr.font.name      = DEFAULT_FONT
            r_ltr.font.size      = Pt(FONT_BODY)
            r_ltr.font.bold      = is_correct
            r_ltr.font.color.rgb = RGB_GREEN if is_correct else RGB_DARK_TEXT

            r_cht = p_ch.add_run(ch_text)
            r_cht.font.name      = DEFAULT_FONT
            r_cht.font.size      = Pt(FONT_BODY)
            r_cht.font.bold      = is_correct
            r_cht.font.color.rgb = RGB_GREEN if is_correct else RGB_DARK_TEXT

        # ── Correct answer banner ─────────────────────────────────────────────
        p_ans = doc.add_paragraph()
        p_ans.paragraph_format.left_indent  = Inches(0.15)
        p_ans.paragraph_format.space_before = Pt(8)
        p_ans.paragraph_format.space_after  = Pt(2)
        add_shading(p_ans, LIGHT_BLUE)
        set_paragraph_format(p_ans, keep_with_next=True)

        r_lbl = p_ans.add_run("✓ Answer: ")
        r_lbl.font.name      = DEFAULT_FONT
        r_lbl.font.size      = Pt(FONT_BODY)
        r_lbl.font.bold      = True
        r_lbl.font.color.rgb = RGB_GREEN

        r_aval = p_ans.add_run(f"{c_letter})  {c_text}")
        r_aval.font.name      = DEFAULT_FONT
        r_aval.font.size      = Pt(FONT_BODY)
        r_aval.font.color.rgb = RGB_DARK_TEXT

        # ── Explanation ───────────────────────────────────────────────────────
        if expl:
            p_expl = doc.add_paragraph()
            p_expl.paragraph_format.left_indent  = Inches(0.15)
            p_expl.paragraph_format.space_before = Pt(4)
            p_expl.paragraph_format.space_after  = Pt(4)
            set_paragraph_format(p_expl)

            r_el = p_expl.add_run("Explanation:  ")
            r_el.font.name      = DEFAULT_FONT
            r_el.font.size      = Pt(FONT_SMALL)
            r_el.font.bold      = True
            r_el.font.color.rgb = RGB_BLUE

            r_et = p_expl.add_run(expl)
            r_et.font.name      = DEFAULT_FONT
            r_et.font.size      = Pt(FONT_SMALL)
            r_et.font.color.rgb = RGB_DARK_TEXT

        # ── Citation ──────────────────────────────────────────────────────────
        if cits:
            p_cit = doc.add_paragraph()
            p_cit.paragraph_format.left_indent  = Inches(0.15)
            p_cit.paragraph_format.space_before = Pt(2)
            p_cit.paragraph_format.space_after  = Pt(8)
            set_paragraph_format(p_cit)

            r_cl = p_cit.add_run("Ref:  ")
            r_cl.font.name      = DEFAULT_FONT
            r_cl.font.size      = Pt(FONT_TINY)
            r_cl.font.bold      = True
            r_cl.font.color.rgb = RGB_GRAY

            r_ct = p_cit.add_run("  |  ".join(cits))
            r_ct.font.name      = DEFAULT_FONT
            r_ct.font.size      = Pt(FONT_TINY)
            r_ct.font.italic    = True
            r_ct.font.color.rgb = RGB_GRAY

        # ── Divider (not after last question) ─────────────────────────────────
        if idx < total_q:
            add_divider(doc)

    doc.save(str(out_path))
    return total_q


# ── XLSX builder ──────────────────────────────────────────────────────────────

def build_xlsx(chunk: list, q_data: dict, out_path: Path) -> int:
    """Write one XLSX for the chunk. Returns question count."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    from word_doc_defaults import sanitize

    HEADER_FILL  = PatternFill("solid", fgColor="1B3564")  # navy
    CORRECT_FILL = PatternFill("solid", fgColor="E6F4EA")  # light green
    ALT_FILL     = PatternFill("solid", fgColor="F5F8FC")  # very light blue

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "AAFP Q&A"

    headers = ["#", "Quiz", "Stem", "A", "B", "C", "D", "E",
               "Correct", "Correct Answer", "Explanation", "Citation"]
    ws.append(headers)

    # Header row styling
    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font      = Font(name="Aptos", bold=True, color="FFFFFF", size=10)
        cell.fill      = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # Column widths
    col_widths = [5, 26, 55, 22, 22, 22, 22, 22, 8, 28, 60, 55]
    for i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # Row height for header
    ws.row_dimensions[1].height = 22

    # Data rows
    qids_ordered = []
    for _, info in chunk:
        qids_ordered.extend(info["qids"])

    for row_idx, aafp_qid in enumerate(qids_ordered, start=1):
        q       = q_data[aafp_qid]
        choices = {ch["letter"]: sanitize(ch.get("text", "")) for ch in q["choices"]}
        cit_str = "  |  ".join(sanitize(c) for c in q["citations"])
        c_ltr   = q["correct_letter"]

        data_row = [
            row_idx,
            q["quiz_title"],
            sanitize(q["stem"]),
            choices.get("A", ""),
            choices.get("B", ""),
            choices.get("C", ""),
            choices.get("D", ""),
            choices.get("E", ""),
            c_ltr,
            sanitize(q["correct_text"]),
            sanitize(q["explanation"]),
            cit_str,
        ]
        ws.append(data_row)

        excel_row = row_idx + 1
        alt = (row_idx % 2 == 0)

        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=excel_row, column=col_idx)
            cell.alignment = Alignment(vertical="top", wrap_text=True)

            if col_idx in (9, 10):   # Correct / Correct Answer
                cell.font = Font(name="Aptos", size=9, bold=True, color="3E8D27")
                cell.fill = CORRECT_FILL
            else:
                cell.font = Font(name="Aptos", size=9)
                if alt:
                    cell.fill = ALT_FILL

    ws.freeze_panes = "A2"
    wb.save(str(out_path))
    return len(qids_ordered)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(
        description="Generate AAFP Board Review Q&A deliverable files (DOCX + XLSX)"
    )
    ap.add_argument("--dry-run",   action="store_true",
                    help="Print chunk plan only — no files written")
    ap.add_argument("--docx-only", action="store_true",
                    help="Write DOCX files only (skip XLSX)")
    ap.add_argument("--xlsx-only", action="store_true",
                    help="Write XLSX files only (skip DOCX)")
    args = ap.parse_args()

    do_docx = not args.xlsx_only
    do_xlsx = not args.docx_only

    print(f"\n{'='*60}")
    print(f"  build_aafp_qa_deliverables.py")
    print(f"  dry-run={args.dry_run} | docx={do_docx} | xlsx={do_xlsx}")
    print(f"{'='*60}")

    # ── Load ─────────────────────────────────────────────────────────────────
    conn = sqlite3.connect(DB_PATH)
    quizzes, q_data = load_all_data(conn)
    conn.close()

    total_q = sum(len(info["qids"]) for _, info in quizzes)
    print(f"  Quizzes: {len(quizzes)}   Questions: {total_q}")

    # ── Chunk ────────────────────────────────────────────────────────────────
    chunks = chunk_quizzes(quizzes)
    print(f"  Files to generate: {len(chunks)}\n")

    for i, chunk in enumerate(chunks, start=1):
        name  = chunk_filename(chunk)
        count = chunk_q_count(chunk)
        first = chunk[0][1]["display_num"]
        last  = chunk[-1][1]["display_num"]
        print(f"  [{i:2d}] {name:<28}  {count:3d} Q   "
              f"(quiz {first:03d}–{last:03d})")

    if args.dry_run:
        print(f"\n  [DRY RUN] No files written.")
        print(f"{'='*60}\n")
        return

    # ── Output dirs ───────────────────────────────────────────────────────────
    OUT_DOCX.mkdir(parents=True, exist_ok=True)
    OUT_XLSX.mkdir(parents=True, exist_ok=True)

    # ── Generate ──────────────────────────────────────────────────────────────
    file_log = []
    for i, chunk in enumerate(chunks, start=1):
        name = chunk_filename(chunk)

        if do_docx:
            docx_path = OUT_DOCX / f"{name}.docx"
            print(f"  [{i:2d}/{len(chunks)}] DOCX: {name} ...", end="", flush=True)
            n = build_docx(chunk, q_data, docx_path)
            print(f"  {n} Q ✓")

        if do_xlsx:
            xlsx_path = OUT_XLSX / f"{name}.xlsx"
            print(f"  [{i:2d}/{len(chunks)}] XLSX: {name} ...", end="", flush=True)
            n = build_xlsx(chunk, q_data, xlsx_path)
            print(f"  {n} Q ✓")

        file_log.append({"file": name, "q_count": chunk_q_count(chunk)})

    # ── Log ───────────────────────────────────────────────────────────────────
    import json as _json
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    ts       = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    log_path = LOG_DIR / f"build_aafp_qa_deliverables_{ts}.json"
    log_path.write_text(_json.dumps({
        "script":    "build_aafp_qa_deliverables.py",
        "timestamp": ts,
        "chunks":    len(chunks),
        "total_q":   total_q,
        "files":     file_log,
    }, indent=2))

    # ── Summary ───────────────────────────────────────────────────────────────
    n_files = len(chunks) * ((1 if do_docx else 0) + (1 if do_xlsx else 0))
    print(f"\n{'='*60}")
    print(f"  Done — {n_files} files written")
    print(f"  DOCX → {OUT_DOCX.relative_to(PROJECT_ROOT)}")
    print(f"  XLSX → {OUT_XLSX.relative_to(PROJECT_ROOT)}")
    print(f"  Log  → {log_path.name}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
