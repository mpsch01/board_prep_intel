#!/usr/bin/env python3
"""
build_custom_question_set.py
============================
Content-addressable practice question set generator.
Queries the ITE + AAFP question banks by blueprint category, body system,
source bank, and/or year range, then produces two DOCX files:
  - Exam version:       questions + answer key table (test-taking mode)
  - Study guide version: questions + correct answer + explanation after each (review mode)

Rule 14 compliant: uses word_doc_defaults.py for all DOCX generation.
Rule 4 compliant:  dynamic paths only.

Usage (from 03_module.3_analyst/scripts/):
    python3 build_custom_question_set.py --count 50 --blueprint "Acute Care and Diagnosis" --bank ITE
    python3 build_custom_question_set.py --count 75 --body-system Cardiovascular --body-system Respiratory --bank BOTH
    python3 build_custom_question_set.py --count 35 --blueprint "Foundations of Care" --body-system Neurologic --bank ITE --years 2020-2025
    python3 build_custom_question_set.py --count 100 --bank BOTH   # fully random

The SKILL parses user language and maps aliases to canonical values before
calling this script. This script only accepts canonical DB values.

Outputs (in custom_question_sets/YYYY-MM-DD/):
    QSet_<N>Q_<label>_Exam.docx
    QSet_<N>Q_<label>_StudyGuide.docx
"""

from __future__ import annotations

import argparse
import json
import random
import re
import sqlite3
import sys
from datetime import date
from pathlib import Path

# ── Path setup (Rule 4) ───────────────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
DB_PATH      = PROJECT_ROOT / "00_database" / "db" / "ite_intelligence.db"
SETS_ROOT    = PROJECT_ROOT / "03_module.3_analyst" / "custom_question_sets"

sys.path.insert(0, str(SCRIPT_DIR))

from word_doc_defaults import (
    new_document, add_title, add_subtitle,
    add_bottom_border, add_left_border, add_shading, set_paragraph_format,
    DEFAULT_FONT, RGB_NAVY, RGB_BLUE, RGB_GRAY, RGB_GOLD,
    NAVY, BLUE, GOLD,
    FONT_TITLE, FONT_SUBTITLE, FONT_HEADING, FONT_BODY, FONT_SMALL, FONT_TINY,
)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Encoding fix ──────────────────────────────────────────────────────────────
_DOTLEADER = re.compile("[" + chr(0xF02E) + chr(0xF02D) + chr(0xF02C) + "]+")

# Symbol-font UTF-8 bytes misread as Windows-1252, and double-encoded Latin chars.
# Pattern: Symbol font char at code point 0xF0XX was stored as UTF-8 (3 bytes: EF 8X XX),
# then those bytes were interpreted individually as Latin-1, producing 3-char sequences.
_ENCODING_FIXES = [
    ("ï‚£", "≤"),   # ï‚£ → ≤  (Symbol 0xA3)
    ("ï‚³", "≥"),   # ï‚³ → ≥  (Symbol 0xB3)
    ("ï‚±", "±"),   # ï‚± → ±  (Symbol 0xB1)
    ("ï‚®", "→"),   # ï‚® → →  (Symbol 0xAE arrow right)
    ("ï‚¬", "←"),   # ï‚¬ → ←  (Symbol 0xAC arrow left)
    ("ï‚´", "×"),   # ï‚´ → ×  (Symbol 0xB4 multiply)
    ("ï‚¸", "÷"),   # ï‚¸ → ÷  (Symbol 0xB8 divide)
    # Double-encoded UTF-8 (Latin chars stored as W-1252 codepoints)
    ("Ã©", "é"),         # Ã© → é
    ("Ã¨", "è"),         # Ã¨ → è
    ("Ã¼", "ü"),         # Ã¼ → ü
    ("Ã¶", "ö"),         # Ã¶ → ö
    ("Â²", "²"),         # Â² → ²
    ("Â³", "³"),         # Â³ → ³
    ("Â°", "°"),         # Â° → °
]

def clean_text(text: str | None) -> str:
    if not text:
        return text or ""
    for bad, good in _ENCODING_FIXES:
        if bad in text:
            text = text.replace(bad, good)
    text = _DOTLEADER.sub(" ", text)
    text = text.replace("“", '"').replace("”", '"')
    text = text.replace("‘", "'").replace("’", "'")
    return text

def clean_choices(raw) -> dict:
    if isinstance(raw, str):
        try:
            parsed = json.loads(raw)
        except Exception:
            return {}
        if isinstance(parsed, list):
            return {c["letter"]: clean_text(c.get("text", "")) for c in parsed if "letter" in c}
        if isinstance(parsed, dict):
            return {k: clean_text(v) for k, v in parsed.items()}
        return {}
    if isinstance(raw, list):
        return {c["letter"]: clean_text(c.get("text", "")) for c in raw if "letter" in c}
    if isinstance(raw, dict):
        return {k: clean_text(v) for k, v in raw.items()}
    return {}


# ── Reference parsing ─────────────────────────────────────────────────────────
_TRAILING_PAGENUM = re.compile(r"[\n\r]+\d{1,3}\s*$")
_REF_SPLIT_NUM    = re.compile(r"\s+\d+\)\s+")  # "  2)  " pattern
_REF_PREFIX       = re.compile(r"^Ref:\s*", re.IGNORECASE)

def split_explanation_and_refs(text: str | None) -> tuple[str, str]:
    """
    Split ITE explanation text into (clinical_body, raw_reference_block).
    References are embedded at the end of the explanation field, starting
    with "Ref:" or a numbered citation pattern.
    Returns (body, refs) — refs is empty string if none found.
    """
    if not text:
        return (text or ""), ""
    text = clean_text(text)
    # Strip stray trailing page numbers (e.g. "\n31" at end of field)
    text = _TRAILING_PAGENUM.sub("", text).strip()
    # Split at "Ref:" marker
    idx = text.find("\nRef:")
    if idx == -1:
        idx = text.find("\nref:")
    if idx >= 0:
        return text[:idx].strip(), text[idx+1:].strip()
    # Fallback: no Ref: found — entire text is body
    return text.strip(), ""

def parse_references(raw_refs: str) -> list[str]:
    """
    Parse a raw reference block into individual citation strings.
    Handles both pipe-separated (|) and numbered (2) 3)) formats.
    """
    if not raw_refs:
        return []
    text = _REF_PREFIX.sub("", raw_refs).strip()
    # Pipe-separated (newer questions)
    if " | " in text or text.count("|") >= 1:
        refs = [r.strip() for r in text.split("|") if r.strip()]
    else:
        # Numbered: split on "  2) " "  3) " etc.
        parts = _REF_SPLIT_NUM.split(text)
        refs  = [p.strip() for p in parts if p.strip()]
    return refs

# ── CLI args ──────────────────────────────────────────────────────────────────
def parse_args():
    p = argparse.ArgumentParser(
        description="Generate a content-filtered practice question set (Exam + Study Guide)."
    )
    p.add_argument("--count",       type=int, required=True,
                   help="Number of questions to draw")
    p.add_argument("--blueprint",   action="append", default=[],
                   help="Blueprint category filter (canonical DB value). Repeatable → OR within blueprint.")
    p.add_argument("--body-system", action="append", default=[],
                   dest="body_system",
                   help="Body system filter (canonical DB value). Repeatable → OR within body_system.")
    p.add_argument("--bank",        default="BOTH",
                   choices=["ITE", "AAFP", "BOTH"],
                   help="Question bank source (default: BOTH)")
    p.add_argument("--years",       default=None,
                   help="Year range for ITE questions: YYYY or YYYY-YYYY (e.g. 2020-2025)")
    p.add_argument("--label",       default=None,
                   help="Short label for output filename (auto-generated if omitted)")
    p.add_argument("--output-dir",  default=None,
                   help="Output directory (defaults to custom_question_sets/YYYY-MM-DD/)")
    p.add_argument("--seed",        type=int, default=None,
                   help="Random seed for reproducibility (default: random each run)")
    return p.parse_args()


# ── Year range parser ─────────────────────────────────────────────────────────
def parse_years(years_str: str | None) -> tuple[int, int] | None:
    """Parse '2022' or '2020-2025' into (start, end) tuple, or None."""
    if not years_str:
        return None
    years_str = years_str.strip()
    m = re.match(r'^(\d{4})(?:-(\d{4}))?$', years_str)
    if not m:
        raise ValueError(f"Invalid year format: {years_str!r}. Use YYYY or YYYY-YYYY.")
    start = int(m.group(1))
    end   = int(m.group(2)) if m.group(2) else start
    return (start, end)


# ── SQL query builder ─────────────────────────────────────────────────────────
def build_query(bank: str, blueprints: list, body_systems: list,
                year_range: tuple | None) -> tuple[str, list]:
    """
    Build a UNION SQL query across ITE (questions) and/or AAFP (aafp_questions).

    Logic:
      - Multiple blueprints    → OR within blueprint  (blueprint IN (...))
      - Multiple body systems  → OR within body_system (body_system IN (...))
      - Blueprint + body_system together → AND across dimensions
      - Year range applies to ITE only (aafp_questions has no exam_year)

    Returns (sql_string, params_list).
    """
    def _where_clauses(table: str, include_years: bool) -> tuple[str, list]:
        clauses = []
        params  = []

        if blueprints:
            placeholders = ",".join("?" * len(blueprints))
            clauses.append(f"blueprint IN ({placeholders})")
            params.extend(blueprints)

        if body_systems:
            placeholders = ",".join("?" * len(body_systems))
            clauses.append(f"body_system IN ({placeholders})")
            params.extend(body_systems)

        if include_years and year_range:
            clauses.append("exam_year BETWEEN ? AND ?")
            params.extend([year_range[0], year_range[1]])

        where = ("WHERE " + " AND ".join(clauses)) if clauses else ""
        return where, params

    ite_cols  = "qid, question_text, choices, correct_letter, correct_text, explanation, blueprint, body_system, exam_year, 'ITE' as source_bank"
    aafp_cols = "aafp_qid as qid, stem as question_text, choices, correct_letter, correct_text, explanation, blueprint, body_system, NULL as exam_year, 'AAFP' as source_bank"

    parts  = []
    params = []

    if bank in ("ITE", "BOTH"):
        where, p = _where_clauses("questions", include_years=True)
        parts.append(f"SELECT {ite_cols} FROM questions {where}")
        params.extend(p)

    if bank in ("AAFP", "BOTH"):
        where, p = _where_clauses("aafp_questions", include_years=False)
        parts.append(f"SELECT {aafp_cols} FROM aafp_questions {where}")
        params.extend(p)

    union_sql = " UNION ALL ".join(parts)
    final_sql = f"SELECT * FROM ({union_sql}) ORDER BY RANDOM()"

    return final_sql, params


# ── DB query ──────────────────────────────────────────────────────────────────
def fetch_questions(db_path: Path, sql: str, params: list, count: int) -> list[dict]:
    """Run the query and return up to count question dicts."""
    con = sqlite3.connect(str(db_path))
    con.row_factory = sqlite3.Row
    cur = con.cursor()
    cur.execute(sql, params)
    rows = cur.fetchmany(count)
    con.close()

    questions = []
    for row in rows:
        d = dict(row)
        d["question_text"] = clean_text(d.get("question_text"))
        d["explanation"]   = clean_text(d.get("explanation"))
        d["choices"]       = clean_choices(d.get("choices"))
        questions.append(d)
    return questions


# ── Label builder ─────────────────────────────────────────────────────────────
def build_label(blueprints: list, body_systems: list, bank: str, year_range) -> str:
    """Short filesystem-safe label from filter params."""
    parts = []

    bp_short = {
        "Acute Care and Diagnosis":  "AcuteCare",
        "Chronic Care Management":   "ChronicCare",
        "Emergent and Urgent Care":  "Emergent",
        "Preventive Care":           "Preventive",
        "Foundations of Care":       "Foundations",
    }
    bs_short = {
        "Cardiovascular":          "Cardio",
        "Respiratory":             "Resp",
        "Injuries/Musculoskeletal":"MSK",
        "Gastrointestinal":        "GI",
        "Endocrine":               "Endo",
        "Psychiatric/Behavioral":  "Psych",
        "Population-Based Care":   "PopHealth",
        "Sexual and Reproductive": "SexRepro",
        "Integumentary":           "Derm",
        "Neurologic":              "Neuro",
        "Nephrologic":             "Nephro",
        "Hematologic/Immune":      "HemeImmune",
        "Special Sensory":         "SpecSensory",
        "Nonspecific":             "Nonspecific",
        "Patient-Based Systems":   "PatientSystems",
    }

    for bp in blueprints:
        parts.append(bp_short.get(bp, bp.replace(" ", "")))
    for bs in body_systems:
        parts.append(bs_short.get(bs, bs.replace("/", "_").replace(" ", "")))

    if not parts:
        parts.append("Random")

    parts.append(bank)

    if year_range:
        if year_range[0] == year_range[1]:
            parts.append(str(year_range[0]))
        else:
            parts.append(f"{year_range[0]}-{year_range[1]}")

    return "_".join(parts)


# ── DOCX helpers ──────────────────────────────────────────────────────────────
def add_page_header(doc: Document, text: str):
    section = doc.sections[0]
    header  = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name      = DEFAULT_FONT
    run.font.size      = Pt(FONT_TINY)
    run.font.color.rgb = RGB_GRAY


def _set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def add_answer_key_table(doc: Document, answer_key: list):
    col_headers = ["#", "QID", "Answer", "Blueprint", "Body System", "Bank"]
    col_widths  = [Inches(0.35), Inches(1.4), Inches(0.6), Inches(1.6), Inches(1.4), Inches(0.65)]

    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, (hdr, w) in enumerate(zip(col_headers, col_widths)):
        hdr_cells[i].width = w
        p   = hdr_cells[i].paragraphs[0]
        run = p.add_run(hdr)
        run.font.name      = DEFAULT_FONT
        run.font.size      = Pt(FONT_SMALL)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr_cells[i], NAVY)

    for row_idx, entry in enumerate(answer_key):
        cells = table.add_row().cells
        vals  = [
            str(entry["num"]),
            entry.get("qid", ""),
            entry.get("letter", ""),
            entry.get("blueprint", ""),
            entry.get("body_system", ""),
            entry.get("bank", ""),
        ]
        fill = "EBF0F7" if row_idx % 2 == 0 else "FFFFFF"
        for i, (val, w) in enumerate(zip(vals, col_widths)):
            cells[i].width = w
            p   = cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(FONT_SMALL)
            _set_cell_shading(cells[i], fill)


def _build_cover(doc: Document, title: str, subtitle_line: str, info_line: str):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(72)
    add_title(doc, title)
    add_subtitle(doc, subtitle_line)
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(6)
    r = info.add_run(info_line)
    r.font.name      = DEFAULT_FONT
    r.font.size      = Pt(FONT_SMALL)
    r.font.color.rgb = RGB_GRAY
    doc.add_page_break()


def _write_question_stem(doc: Document, num: int, q_text: str, choices: dict):
    """Shared: Question N header + stem + choices."""
    qh = doc.add_paragraph()
    qh.paragraph_format.space_before = Pt(16)
    qh.paragraph_format.space_after  = Pt(4)
    set_paragraph_format(qh, keep_with_next=True)
    rn = qh.add_run(f"Question {num}")
    rn.font.name      = DEFAULT_FONT
    rn.font.size      = Pt(FONT_HEADING)
    rn.font.bold      = True
    rn.font.color.rgb = RGB_NAVY

    qs = doc.add_paragraph()
    qs.paragraph_format.space_before = Pt(2)
    qs.paragraph_format.space_after  = Pt(6)
    set_paragraph_format(qs, keep_with_next=True)
    rt = qs.add_run(q_text)
    rt.font.name = DEFAULT_FONT
    rt.font.size = Pt(FONT_BODY)

    if isinstance(choices, dict):
        choice_items = sorted(choices.items())
    elif isinstance(choices, list):
        choice_items = [(c.get("letter", ""), c.get("text", "")) for c in choices]
    else:
        choice_items = []

    for letter, choice_text in choice_items:
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent  = Inches(0.35)
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(2)
        rc = cp.add_run(f"{letter}. {clean_text(choice_text)}")
        rc.font.name = DEFAULT_FONT
        rc.font.size = Pt(FONT_BODY)


def build_exam_docx(questions: list, output_path: Path, label: str, filter_summary: str):
    """Exam version: questions + answer key table at end."""
    doc = new_document()
    add_page_header(doc, f"Practice Question Set — Exam Version")

    _build_cover(
        doc,
        title        = "PRACTICE QUESTION SET",
        subtitle_line= f"Exam Version  |  {filter_summary}",
        info_line    = f"{len(questions)} Questions  |  Answer key at end  |  St. Luke's Family Medicine Residency",
    )

    answer_key = []
    for i, q in enumerate(questions, 1):
        _write_question_stem(doc, i, q.get("question_text", ""), q.get("choices", {}))

        answer_key.append({
            "num":        i,
            "qid":        q.get("qid", ""),
            "letter":     q.get("correct_letter", ""),
            "blueprint":  q.get("blueprint", ""),
            "body_system":q.get("body_system", ""),
            "bank":       q.get("source_bank", ""),
        })

        if i < len(questions):
            div = doc.add_paragraph()
            div.paragraph_format.space_before = Pt(10)
            div.paragraph_format.space_after  = Pt(0)
            add_bottom_border(div, color="CCCCCC", size=4)

    doc.add_page_break()
    ak_hdr = doc.add_paragraph()
    ak_hdr.paragraph_format.space_before = Pt(0)
    ak_hdr.paragraph_format.space_after  = Pt(12)
    add_left_border(ak_hdr, color=GOLD, size=24)
    add_shading(ak_hdr, "EFF3FA")
    set_paragraph_format(ak_hdr, keep_with_next=True)
    rk = ak_hdr.add_run("ANSWER KEY")
    rk.font.name      = DEFAULT_FONT
    rk.font.size      = Pt(FONT_HEADING)
    rk.font.bold      = True
    rk.font.color.rgb = RGB_NAVY

    add_answer_key_table(doc, answer_key)
    doc.save(str(output_path))


def build_study_guide_docx(questions: list, output_path: Path, label: str, filter_summary: str):
    """Study guide version: questions + correct answer + explanation after each."""
    doc = new_document()
    add_page_header(doc, f"Practice Question Set — Study Guide")

    _build_cover(
        doc,
        title        = "PRACTICE QUESTION SET",
        subtitle_line= f"Study Guide Version  |  {filter_summary}",
        info_line    = f"{len(questions)} Questions with Explanations  |  St. Luke's Family Medicine Residency",
    )

    for i, q in enumerate(questions, 1):
        _write_question_stem(doc, i, q.get("question_text", ""), q.get("choices", {}))

        # Correct answer line — navy bold
        correct_letter = q.get("correct_letter", "")
        correct_text   = clean_text(q.get("correct_text", ""))
        ans_p = doc.add_paragraph()
        ans_p.paragraph_format.space_before = Pt(8)
        ans_p.paragraph_format.space_after  = Pt(4)
        ans_run = ans_p.add_run(f"✓  Correct Answer: {correct_letter}.  {correct_text}")
        ans_run.font.name      = DEFAULT_FONT
        ans_run.font.size      = Pt(FONT_BODY)
        ans_run.font.bold      = True
        ans_run.font.color.rgb = RGB_NAVY

        # Explanation — split body from references, then render
        explanation_raw = q.get("explanation") or ""
        body_text, raw_refs = split_explanation_and_refs(explanation_raw)
        if not body_text:
            body_text = "No explanation available."

        # Explanation body — shaded box with left border
        exp_p = doc.add_paragraph()
        exp_p.paragraph_format.space_before = Pt(2)
        exp_p.paragraph_format.space_after  = Pt(4)
        exp_p.paragraph_format.left_indent  = Inches(0.2)
        add_shading(exp_p, "EFF3FA")
        add_left_border(exp_p, color=BLUE, size=12)
        exp_run = exp_p.add_run(body_text)
        exp_run.font.name = DEFAULT_FONT
        exp_run.font.size = Pt(FONT_SMALL)

        # References section — each citation as its own paragraph
        refs = parse_references(raw_refs)
        if refs:
            ref_hdr_p = doc.add_paragraph()
            ref_hdr_p.paragraph_format.space_before = Pt(6)
            ref_hdr_p.paragraph_format.space_after  = Pt(2)
            ref_hdr_p.paragraph_format.left_indent  = Inches(0.2)
            ref_hdr_run = ref_hdr_p.add_run("References")
            ref_hdr_run.font.name      = DEFAULT_FONT
            ref_hdr_run.font.size      = Pt(FONT_SMALL)
            ref_hdr_run.font.bold      = True
            ref_hdr_run.font.color.rgb = RGB_NAVY

            for ref_text in refs:
                ref_p = doc.add_paragraph()
                ref_p.paragraph_format.space_before = Pt(2)
                ref_p.paragraph_format.space_after  = Pt(3)
                ref_p.paragraph_format.left_indent  = Inches(0.2)
                ref_run = ref_p.add_run(ref_text)
                ref_run.font.name      = DEFAULT_FONT
                ref_run.font.size      = Pt(FONT_TINY)
                ref_run.font.color.rgb = RGB_GRAY

        # Bank + year tag (right-aligned metadata)
        bank     = q.get("source_bank", "")
        yr       = q.get("exam_year")
        yr_label = f" · {yr}" if yr else ""
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta_p.paragraph_format.space_before = Pt(2)
        meta_p.paragraph_format.space_after  = Pt(4)
        qid_label = q.get("qid", "")
        meta_run = meta_p.add_run(f"{bank}{yr_label}  |  {q.get('blueprint','')}  |  {q.get('body_system','')}  |  {qid_label}")
        meta_run.font.name      = DEFAULT_FONT
        meta_run.font.size      = Pt(FONT_TINY)
        meta_run.font.color.rgb = RGB_GRAY

        if i < len(questions):
            div = doc.add_paragraph()
            div.paragraph_format.space_before = Pt(10)
            div.paragraph_format.space_after  = Pt(0)
            add_bottom_border(div, color="CCCCCC", size=4)

    doc.save(str(output_path))


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════
def main():
    args = parse_args()

    # Resolve seed before anything else so ORDER BY RANDOM() isn't affected
    # (SQLite handles its own RANDOM(); Python seed only affects our shuffle)
    if args.seed is not None:
        random.seed(args.seed)

    year_range = parse_years(args.years)

    # ── Output directory ──────────────────────────────────────────────────
    if args.output_dir:
        output_dir = Path(args.output_dir).resolve()
    else:
        output_dir = SETS_ROOT / date.today().isoformat()
    output_dir.mkdir(parents=True, exist_ok=True)

    # ── Label for filenames ───────────────────────────────────────────────
    label = args.label or build_label(args.blueprint, args.body_system, args.bank, year_range)

    # ── Filter summary for cover pages ───────────────────────────────────
    parts = []
    if args.blueprint:
        parts.append("Blueprint: " + " + ".join(args.blueprint))
    if args.body_system:
        parts.append("System: " + " + ".join(args.body_system))
    if not args.blueprint and not args.body_system:
        parts.append("All categories (random)")
    parts.append(f"Bank: {args.bank}")
    if year_range:
        yr_str = str(year_range[0]) if year_range[0] == year_range[1] else f"{year_range[0]}–{year_range[1]}"
        parts.append(f"Years: {yr_str}")
    filter_summary = "  |  ".join(parts)

    print("=" * 60)
    print("Custom Question Set Generator")
    print("=" * 60)
    print(f"  Count:       {args.count}")
    print(f"  Blueprint:   {args.blueprint or '(any)'}")
    print(f"  Body system: {args.body_system or '(any)'}")
    print(f"  Bank:        {args.bank}")
    print(f"  Years:       {args.years or '(all)'}")
    print(f"  Label:       {label}")
    print(f"  Output dir:  {output_dir}")
    print()

    # ── Build and run query ───────────────────────────────────────────────
    sql, params = build_query(args.bank, args.blueprint, args.body_system, year_range)
    print("STEP 1: Query database")
    questions = fetch_questions(DB_PATH, sql, params, args.count)

    pool_size = len(questions)
    print(f"  Candidates drawn: {pool_size} / {args.count} requested")

    if pool_size == 0:
        print("\n  ✗ No questions found matching the specified filters.")
        print("  Check filter values or broaden the query.")
        sys.exit(1)

    if pool_size < args.count:
        print(f"  ⚠ Pool smaller than requested — producing {pool_size} questions")

    from collections import Counter
    bp_dist  = Counter(q.get("blueprint") for q in questions)
    bs_dist  = Counter(q.get("body_system") for q in questions)
    bk_dist  = Counter(q.get("source_bank") for q in questions)

    print(f"  Bank split:  " + "  ".join(f"{k}:{v}" for k, v in sorted(bk_dist.items())))
    print(f"  Blueprint:   " + "  ".join(f"{k}:{v}" for k, v in sorted(bp_dist.items())))
    print()

    # ── Generate DOCX files ───────────────────────────────────────────────
    print("STEP 2: Generate DOCX files")

    exam_path  = output_dir / f"QSet_{pool_size}Q_{label}_Exam.docx"
    guide_path = output_dir / f"QSet_{pool_size}Q_{label}_StudyGuide.docx"

    print(f"  Building Exam version...", end=" ", flush=True)
    build_exam_docx(questions, exam_path, label, filter_summary)
    print(f"✓  {exam_path.name}")

    print(f"  Building Study Guide version...", end=" ", flush=True)
    build_study_guide_docx(questions, guide_path, label, filter_summary)
    print(f"✓  {guide_path.name}")

    print()
    print("=" * 60)
    print("COMPLETE")
    print("=" * 60)
    print(f"\n2 files written to:")
    print(f"  {output_dir}")
    print(f"    {exam_path.name}")
    print(f"    {guide_path.name}")


if __name__ == "__main__":
    main()
