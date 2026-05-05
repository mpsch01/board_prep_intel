#!/usr/bin/env python3
"""
build_exam_series.py
====================
Generalized practice exam series generator for any resident with a completed
ITE analysis JSON.  Produces N exam files, each with M unique questions,
weakness-weighted by the resident's ITE performance profile and randomly
distributed across files.

Also merges any practice questions already embedded in the analysis report so
the resident gets all relevant material across the exam series.

Rule 14 compliant: uses word_doc_defaults.py for all DOCX generation.
Rule 4 compliant: dynamic paths only.

Usage (from 03_module.3_analyst/scripts/):
    python3 build_exam_series.py \\
        --resident-dir   "../resident_data/ITE_okezia_cole" \\
        --resident-name  "Okezia Cole, D.O." \\
        --pgy            1 \\
        --num-exams      4 \\
        --questions      50

All arguments have defaults (4 exams × 50 questions). The script auto-discovers
the most recent analysis_v*_*.json in the resident's outputs/ folder if
--analysis is not supplied.

Outputs (all in <resident-dir>/outputs/):
    ITE_<YEAR>_Practice_Exam_<A|B|…>_<LastName>.docx   (N files)
"""

from __future__ import annotations

import argparse
import json
import random
import re
import sys
from collections import Counter
from pathlib import Path

# ── Path setup (Rule 4) ───────────────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
DB_PATH      = PROJECT_ROOT / "00_database" / "db" / "ite_intelligence.db"

sys.path.insert(0, str(SCRIPT_DIR))

# Rule 14: import word_doc_defaults
from word_doc_defaults import (
    new_document, add_title, add_subtitle,
    add_bottom_border, add_left_border, add_shading, set_paragraph_format,
    DEFAULT_FONT, RGB_NAVY, RGB_BLUE, RGB_GRAY, RGB_GOLD,
    NAVY, BLUE, GOLD,
    FONT_TITLE, FONT_SUBTITLE, FONT_HEADING, FONT_BODY, FONT_SMALL, FONT_TINY,
)
from ite_analyzer_v3 import match_practice_questions_v3

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── CLI args ──────────────────────────────────────────────────────────────────
def parse_args():
    p = argparse.ArgumentParser(description="Generate practice exam series from ITE analysis JSON.")
    p.add_argument("--resident-dir",  required=True,
                   help="Path to resident directory (e.g. ../resident_data/ITE_okezia_cole)")
    p.add_argument("--analysis",      default=None,
                   help="Path to analysis JSON (auto-discovered if omitted)")
    p.add_argument("--resident-name", default=None,
                   help="Display name for cover page (e.g. 'Okezia Cole, D.O.'). "
                        "Auto-derived from directory name if omitted.")
    p.add_argument("--pgy",           type=int, default=None,
                   help="PGY level (e.g. 1). Omit to leave blank on cover.")
    p.add_argument("--num-exams",     type=int, default=4,
                   help="Number of exam files to generate (default: 4)")
    p.add_argument("--questions",     type=int, default=50,
                   help="Questions per exam (default: 50)")
    p.add_argument("--seed",          type=int, default=42,
                   help="Random seed for reproducible distribution (default: 42)")
    return p.parse_args()


# ── Encoding fix ──────────────────────────────────────────────────────────────
# U+F02E = Symbol-font private-use period (dot-leader in original ITE PDFs).
# Renders as garbage (ï€®) in Word without Symbol font embedded.
_DOTLEADER = re.compile(r'[]+')

def clean_text(text: str | None) -> str:
    """Replace Symbol-font dot-leader chars and normalize Unicode punctuation."""
    if not text:
        return text or ""
    text = _DOTLEADER.sub(" ", text)
    text = text.replace("“", '"').replace("”", '"')
    text = text.replace("‘", "'").replace("’", "'")
    return text

def clean_choices(raw) -> dict:
    """Normalize choices to {letter: text} dict regardless of source format."""
    if isinstance(raw, str):
        try:
            parsed = json.loads(raw)
        except Exception:
            return {}
        if isinstance(parsed, list):
            return {c["letter"]: clean_text(c.get("text", "")) for c in parsed if "letter" in c}
        if isinstance(parsed, dict):
            return {k: clean_text(v) for k, v in parsed.items()}
        return {}
    if isinstance(raw, list):
        return {c["letter"]: clean_text(c.get("text", "")) for c in raw if "letter" in c}
    if isinstance(raw, dict):
        return {k: clean_text(v) for k, v in raw.items()}
    return {}


# ── Auto-discovery ────────────────────────────────────────────────────────────
def find_analysis_json(resident_dir: Path) -> Path:
    """Return the most recently modified analysis_v*_*.json in outputs/."""
    candidates = sorted(
        (resident_dir / "outputs").glob("analysis_v*.json"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if not candidates:
        raise FileNotFoundError(
            f"No analysis_v*.json found in {resident_dir / 'outputs'}"
        )
    return candidates[0]

def derive_display_name(resident_dir: Path) -> str:
    """Turn 'ITE_okezia_cole' → 'Okezia Cole'."""
    parts = resident_dir.name.split("_")
    # Strip leading 'ITE' or 'AAFP' prefix if present
    if parts and parts[0].upper() in ("ITE", "AAFP"):
        parts = parts[1:]
    return " ".join(p.capitalize() for p in parts)

def derive_last_name(resident_dir: Path) -> str:
    """Turn 'ITE_okezia_cole' → 'Cole' (used in output filenames)."""
    parts = resident_dir.name.split("_")
    if parts and parts[0].upper() in ("ITE", "AAFP"):
        parts = parts[1:]
    return parts[-1].capitalize() if parts else "Resident"


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX helpers
# ═══════════════════════════════════════════════════════════════════════════════

def add_page_header(doc: Document, text: str):
    """Add right-aligned running header to all pages."""
    section = doc.sections[0]
    header  = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name      = DEFAULT_FONT
    run.font.size      = Pt(FONT_TINY)
    run.font.color.rgb = RGB_GRAY


def add_answer_key_table(doc: Document, answer_key: list):
    """Render answer key as a formatted table."""
    col_headers = ["#", "QID", "Answer", "Target Area", "Body System"]
    col_widths  = [Inches(0.4), Inches(1.55), Inches(0.65), Inches(1.85), Inches(1.55)]

    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = "Table Grid"

    # Header row — navy background, white text
    hdr_cells = table.rows[0].cells
    for i, (hdr, w) in enumerate(zip(col_headers, col_widths)):
        hdr_cells[i].width = w
        p   = hdr_cells[i].paragraphs[0]
        run = p.add_run(hdr)
        run.font.name      = DEFAULT_FONT
        run.font.size      = Pt(FONT_SMALL)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  NAVY)
        tcPr.append(shd)

    # Data rows — alternating light blue
    for row_idx, entry in enumerate(answer_key):
        cells = table.add_row().cells
        vals  = [
            str(entry["num"]),
            entry.get("qid", ""),
            entry.get("letter", ""),
            entry.get("target", ""),
            entry.get("body", ""),
        ]
        fill = "EBF0F7" if row_idx % 2 == 0 else "FFFFFF"
        for i, (val, w) in enumerate(zip(vals, col_widths)):
            cells[i].width = w
            p   = cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(FONT_SMALL)
            tcPr = cells[i]._tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tcPr.append(shd)


def build_exam_docx(
    questions:      list,
    exam_letter:    str,
    output_path:    Path,
    resident_name:  str,
    pgy_label:      str,
    exam_year:      str,
):
    """Build a single practice exam DOCX."""
    doc = new_document()   # Rule 14: word_doc_defaults page setup + Aptos font

    add_page_header(doc, f"ITE {exam_year} Practice Exam {exam_letter}")

    # ── Cover page ────────────────────────────────────────────────────────
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(72)

    add_title(doc, f"ITE {exam_year} PRACTICE EXAM {exam_letter}")
    add_subtitle(doc, f"{resident_name}  |  {pgy_label}  |  St. Luke's Family Medicine Residency")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(6)
    r = info.add_run(
        f"{len(questions)} Questions — Weighted by Performance Priority  |  Answer key at end"
    )
    r.font.name      = DEFAULT_FONT
    r.font.size      = Pt(FONT_SMALL)
    r.font.color.rgb = RGB_GRAY

    doc.add_page_break()

    # ── Questions ─────────────────────────────────────────────────────────
    answer_key = []

    for i, q in enumerate(questions, 1):
        q_text         = q.get("question_text", "")
        choices        = q.get("choices") or {}
        targeting      = q.get("targeting", "")
        body_sys       = q.get("body_system_merged") or q.get("body_system", "")
        qid            = q.get("qid", "")
        correct_letter = q.get("correct_letter", "")
        correct_text   = q.get("correct_text", "")

        answer_key.append({
            "num":    i,
            "qid":    qid,
            "letter": correct_letter,
            "text":   correct_text,
            "target": targeting,
            "body":   body_sys,
        })

        # "Question N" label
        qh = doc.add_paragraph()
        qh.paragraph_format.space_before = Pt(16)
        qh.paragraph_format.space_after  = Pt(4)
        set_paragraph_format(qh, keep_with_next=True)
        rn = qh.add_run(f"Question {i}")
        rn.font.name      = DEFAULT_FONT
        rn.font.size      = Pt(FONT_HEADING)
        rn.font.bold      = True
        rn.font.color.rgb = RGB_NAVY

        # Question stem
        qs = doc.add_paragraph()
        qs.paragraph_format.space_before = Pt(2)
        qs.paragraph_format.space_after  = Pt(6)
        set_paragraph_format(qs, keep_with_next=True)
        rt = qs.add_run(q_text)
        rt.font.name = DEFAULT_FONT
        rt.font.size = Pt(FONT_BODY)

        # Answer choices
        if isinstance(choices, dict):
            choice_items = sorted(choices.items())
        elif isinstance(choices, list):
            choice_items = [(c.get("letter", ""), c.get("text", "")) for c in choices]
        else:
            choice_items = []

        for letter, choice_text in choice_items:
            cp = doc.add_paragraph()
            cp.paragraph_format.left_indent  = Inches(0.35)
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after  = Pt(2)
            rc = cp.add_run(f"{letter}. {clean_text(choice_text)}")
            rc.font.name = DEFAULT_FONT
            rc.font.size = Pt(FONT_BODY)

        # Thin divider between questions
        if i < len(questions):
            div = doc.add_paragraph()
            div.paragraph_format.space_before = Pt(10)
            div.paragraph_format.space_after  = Pt(0)
            add_bottom_border(div, color="CCCCCC", size=4)

    # ── Answer key page ───────────────────────────────────────────────────
    doc.add_page_break()

    ak_hdr = doc.add_paragraph()
    ak_hdr.paragraph_format.space_before = Pt(0)
    ak_hdr.paragraph_format.space_after  = Pt(12)
    add_left_border(ak_hdr, color=GOLD, size=24)
    add_shading(ak_hdr, "EFF3FA")
    set_paragraph_format(ak_hdr, keep_with_next=True)
    rk = ak_hdr.add_run("ANSWER KEY")
    rk.font.name      = DEFAULT_FONT
    rk.font.size      = Pt(FONT_HEADING)
    rk.font.bold      = True
    rk.font.color.rgb = RGB_NAVY

    add_answer_key_table(doc, answer_key)

    doc.save(str(output_path))


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════
def main():
    args = parse_args()

    resident_dir = Path(args.resident_dir).resolve()
    if not resident_dir.exists():
        print(f"ERROR: resident directory not found: {resident_dir}")
        sys.exit(1)

    output_dir = resident_dir / "outputs"
    output_dir.mkdir(parents=True, exist_ok=True)

    # ── Resolve analysis JSON ─────────────────────────────────────────────
    if args.analysis:
        analysis_path = Path(args.analysis).resolve()
    else:
        analysis_path = find_analysis_json(resident_dir)

    print("=" * 60)
    print("ITE Practice Exam Series Generator")
    print("=" * 60)
    print(f"  Resident dir:   {resident_dir}")
    print(f"  Analysis JSON:  {analysis_path}")
    print(f"  Exams:          {args.num_exams}")
    print(f"  Q per exam:     {args.questions}")
    print(f"  Total pool:     {args.num_exams * args.questions}")
    print(f"  Seed:           {args.seed}")
    print()

    # ── Derive display values ─────────────────────────────────────────────
    resident_name = args.resident_name or derive_display_name(resident_dir)
    last_name     = derive_last_name(resident_dir)
    pgy_label     = f"PGY-{args.pgy}" if args.pgy else "Resident"

    # ── STEP 1: Load and clean analysis JSON ─────────────────────────────
    print("STEP 1: Load and clean analysis JSON")
    with open(analysis_path, encoding="utf-8") as f:
        analysis = json.load(f)

    mid = analysis.get("missed_items_detail", [])
    for m in mid:
        if m.get("question_text"):
            m["question_text"] = clean_text(m["question_text"])
        if m.get("explanation"):
            m["explanation"]   = clean_text(m["explanation"])

    for q in analysis.get("practice_questions", []):
        if q.get("question_text"):
            q["question_text"] = clean_text(q["question_text"])
        if q.get("explanation"):
            q["explanation"]   = clean_text(q["explanation"])

    with open(analysis_path, "w", encoding="utf-8") as f:
        json.dump(analysis, f, indent=2, ensure_ascii=False)

    exam_year = str(analysis.get("exam_year", "20XX"))
    print(f"  ✓ Analysis loaded — exam year: {exam_year}")
    print()

    # ── STEP 2: Build weakness profile ───────────────────────────────────
    print("STEP 2: Build weakness profile")

    items = [
        {
            "item":        m["item_number"],
            "correct":     False,
            "blueprint":   m.get("blueprint"),
            "body_system": m.get("body_system"),
            "score":       500,
        }
        for m in mid if m.get("item_number")
    ]

    qid_map = {
        m["item_number"]: m["qid"]
        for m in mid if m.get("item_number") and m.get("qid")
    }

    icd10_clusters = analysis.get("icd10_weakness_map", {}).get("icd10_clusters", [])
    icd10_profile  = {
        c["code"]: c["miss_count"]
        for c in icd10_clusters
        if c.get("code") and c.get("miss_count")
    }

    priorities = analysis.get("yield_priorities", [])
    concepts   = analysis.get("concept_clustering", {})
    perf       = analysis.get("performance", {})

    total_needed = args.num_exams * args.questions
    print(f"  Missed items:  {len(items)}")
    print(f"  Weak dims:     {len(priorities)}")
    print(f"  ICD-10 codes:  {len(icd10_profile)}")
    print(f"  Target pool:   {total_needed}")
    print()

    # ── STEP 3: Pull weakness-weighted questions ──────────────────────────
    print("STEP 3: Pull questions from DB")

    raw_pool = match_practice_questions_v3(
        perf=perf,
        priorities=priorities,
        qid_map=qid_map,
        items=items,
        db_path=str(DB_PATH),
        target_count=total_needed,
        concepts=concepts,
        current_exam_year=exam_year,
        icd10_profile=icd10_profile,
    )

    print(f"  Returned:  {len(raw_pool)} questions")
    tier_counts = Counter(q.get("match_tier") for q in raw_pool)
    bank_counts = Counter(q.get("source_bank") for q in raw_pool)
    print(f"  Tiers:     " + "  ".join(f"T{k}:{v}" for k, v in sorted(tier_counts.items())))
    print(f"  Banks:     " + "  ".join(f"{k}:{v}" for k, v in sorted(bank_counts.items())))

    # Apply encoding fix
    for q in raw_pool:
        q["question_text"] = clean_text(q.get("question_text", ""))
        q["explanation"]   = clean_text(q.get("explanation", ""))
        q["choices"]       = clean_choices(q.get("choices"))

    # ── STEP 4: Merge original report questions ───────────────────────────
    print()
    print("STEP 4: Merge original report questions")

    raw_pool_qids   = {q.get("qid") for q in raw_pool if q.get("qid")}
    original_qs     = analysis.get("practice_questions", [])
    missing_from_pool = [
        q for q in original_qs
        if q.get("qid") and q.get("qid") not in raw_pool_qids
    ]

    n_inject = len(missing_from_pool)
    if n_inject == 0:
        print(f"  ✓ All {len(original_qs)} report questions already in pool")
        final_pool = raw_pool
    else:
        print(f"  ✓ {n_inject} report question(s) not in ranked pool")
        print(f"    Injecting {n_inject} | Clipping {n_inject} least-fitting from bottom")
        for q in missing_from_pool:
            q["question_text"] = clean_text(q.get("question_text", ""))
            q["explanation"]   = clean_text(q.get("explanation", ""))
            q["choices"]       = clean_choices(q.get("choices"))
        # Keep top (total_needed - n_inject) from ranked pool + inject originals
        clip_at = max(0, len(raw_pool) - n_inject)
        final_pool = raw_pool[:clip_at] + missing_from_pool

    actual_total = len(final_pool)
    print(f"  Final pool: {actual_total} questions")

    if actual_total < total_needed:
        print(f"  ⚠ Pool is smaller than requested ({actual_total} < {total_needed})")

    # ── STEP 5: Distribute across exams ──────────────────────────────────
    print()
    print("STEP 5: Distribute questions")

    per_exam = min(args.questions, actual_total // args.num_exams)

    random.seed(args.seed)
    shuffled = final_pool[:]
    random.shuffle(shuffled)

    exam_letters = [chr(ord("A") + i) for i in range(args.num_exams)]
    exams = {
        letter: shuffled[i * per_exam : (i + 1) * per_exam]
        for i, letter in enumerate(exam_letters)
    }

    for letter, qs in exams.items():
        t = Counter(q.get("match_tier") for q in qs)
        b = Counter(q.get("source_bank") for q in qs)
        print(f"  Exam {letter}: {len(qs)} q | "
              f"T1:{t.get(1,0)} T2:{t.get(2,0)} T3:{t.get(3,0)} | "
              f"ITE:{b.get('ITE',0)} AAFP:{b.get('AAFP',0)}")

    # ── STEP 6: Generate DOCX files ───────────────────────────────────────
    print()
    print("STEP 6: Generate DOCX files")
    print()

    generated = []
    for letter, qs in exams.items():
        fname = f"ITE_{exam_year}_Practice_Exam_{letter}_{last_name}.docx"
        out   = output_dir / fname
        print(f"  Building Exam {letter} ({len(qs)} questions)...", end=" ", flush=True)
        build_exam_docx(
            questions=qs,
            exam_letter=letter,
            output_path=out,
            resident_name=resident_name,
            pgy_label=pgy_label,
            exam_year=exam_year,
        )
        generated.append(out)
        print(f"✓  {fname}")

    print()
    print("=" * 60)
    print("COMPLETE")
    print("=" * 60)
    print(f"\n{len(generated)} exam file(s) written to:")
    print(f"  {output_dir}")
    for p in generated:
        print(f"    {p.name}")


if __name__ == "__main__":
    main()
