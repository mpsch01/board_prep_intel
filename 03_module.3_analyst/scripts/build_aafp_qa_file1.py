"""
build_aafp_qa_file1.py
Builds a Q&A Word document from AAFP questions that are semantic near-duplicates
of ITE questions (File 1 — dist < 0.30, with linked articles), sorted by distance ASC.
Uses word_doc_defaults.py for St. Luke's / ITE Intelligence styling.

Each question is paired with the matched ITE stem in a companion block below the
explanation for side-by-side review.

Run from anywhere — paths are dynamic.
Output: 03_module.3_analyst/reports/AAFP_BRQ_NearDuplicate_QA.docx
"""

import json
import sqlite3
import sys
from pathlib import Path

# ── Paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
DB_PATH      = PROJECT_ROOT / "00_database" / "db" / "ite_intelligence.db"
DEFAULTS_DIR = SCRIPT_DIR   # word_doc_defaults.py lives in the same directory
OUT_PATH     = PROJECT_ROOT / "03_module.3_analyst" / "reports" / "AAFP_BRQ_NearDuplicate_QA.docx"

sys.path.insert(0, str(DEFAULTS_DIR))
from word_doc_defaults import (
    new_document, add_title, add_subtitle, add_section_header, add_body_text,
    add_divider, add_left_border, add_shading, add_bottom_border, add_page_number_footer,
    set_paragraph_format, sanitize,
    NAVY, GOLD, BLUE, LIGHT_BLUE, DARK_TEXT, MED_GRAY,
    RGB_NAVY, RGB_GOLD, RGB_BLUE, RGB_DARK_TEXT, RGB_GRAY, RGB_GREEN,
    DEFAULT_FONT, FONT_HEADING, FONT_BODY, FONT_SMALL, FONT_TINY,
)

from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Config ────────────────────────────────────────────────────────────────────
DIST_THRESHOLD = 0.30   # include AAFP questions with ite_nearest_dist < this value

# ── Query ─────────────────────────────────────────────────────────────────────
QUERY = """
SELECT
    aq.aafp_qid,
    aq.stem,
    aq.choices,
    aq.correct_letter,
    aq.correct_text,
    aq.explanation,
    aq.quiz_title,
    aq.question_number,
    aq.body_system,
    aq.blueprint,
    aq.url,
    aq.ite_nearest_qid,
    aq.ite_nearest_dist,
    q.question_text   AS ite_stem,
    q.exam_year       AS ite_exam_year,
    q.body_system     AS ite_body_system,
    COUNT(DISTINCT x.article_id) AS linked_articles
FROM aafp_questions aq
JOIN questions q
    ON q.qid = aq.ite_nearest_qid
JOIN aafp_qid_art_xref x
    ON x.aafp_qid = aq.aafp_qid
WHERE aq.ite_nearest_qid IS NOT NULL
  AND aq.ite_nearest_dist IS NOT NULL
  AND aq.ite_nearest_dist < ?
GROUP BY
    aq.aafp_qid, aq.stem, aq.choices, aq.correct_letter, aq.correct_text,
    aq.explanation, aq.quiz_title, aq.question_number, aq.body_system,
    aq.blueprint, aq.url, aq.ite_nearest_qid, aq.ite_nearest_dist,
    q.question_text, q.exam_year, q.body_system
ORDER BY aq.ite_nearest_dist ASC
"""

# ── Helpers ───────────────────────────────────────────────────────────────────

def dist_label(dist):
    if dist < 0.25: return f"Near-Duplicate  (dist: {dist:.3f})"
    if dist < 0.27: return f"Strong Match  (dist: {dist:.3f})"
    return f"Semantic Match  (dist: {dist:.3f})"

def dist_color(dist):
    if dist < 0.25: return RGB_GOLD
    if dist < 0.27: return RGB_BLUE
    return RGB_DARK_TEXT

def add_badge_line(doc, dist, quiz_title, q_number, body_system, ite_qid, ite_year):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.25)
    set_paragraph_format(p, keep_with_next=True)

    badge_run = p.add_run(dist_label(dist))
    badge_run.font.name  = DEFAULT_FONT
    badge_run.font.size  = Pt(FONT_TINY)
    badge_run.font.bold  = True
    badge_run.font.color.rgb = dist_color(dist)

    sep = p.add_run("   |   ")
    sep.font.name  = DEFAULT_FONT
    sep.font.size  = Pt(FONT_TINY)
    sep.font.color.rgb = RGB_GRAY

    info = p.add_run(f"Matched ITE: {ite_qid} ({ite_year})   ·   {quiz_title} Q{q_number}   ·   {body_system or 'N/A'}")
    info.font.name  = DEFAULT_FONT
    info.font.size  = Pt(FONT_TINY)
    info.font.color.rgb = RGB_GRAY
    return p

def add_stem(doc, number, stem_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.25)
    set_paragraph_format(p, keep_with_next=True)

    num_run = p.add_run(f"Question {number}  ")
    num_run.font.name  = DEFAULT_FONT
    num_run.font.size  = Pt(FONT_HEADING)
    num_run.font.bold  = True
    num_run.font.color.rgb = RGB_NAVY

    stem_run = p.add_run(sanitize(stem_text) if stem_text else "")
    stem_run.font.name  = DEFAULT_FONT
    stem_run.font.size  = Pt(FONT_BODY)
    stem_run.font.color.rgb = RGB_DARK_TEXT
    return p

def add_choice(doc, letter, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    set_paragraph_format(p, keep_with_next=True)

    letter_run = p.add_run(f"{letter}.  ")
    letter_run.font.name  = DEFAULT_FONT
    letter_run.font.size  = Pt(FONT_BODY)
    letter_run.font.color.rgb = RGB_DARK_TEXT

    text_run = p.add_run(sanitize(text) if text else "")
    text_run.font.name  = DEFAULT_FONT
    text_run.font.size  = Pt(FONT_BODY)
    text_run.font.color.rgb = RGB_DARK_TEXT
    return p

def add_correct_answer_line(doc, letter, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.25)
    set_paragraph_format(p, keep_with_next=True)

    check_run = p.add_run("CORRECT ANSWER: ")
    check_run.font.name  = DEFAULT_FONT
    check_run.font.size  = Pt(FONT_BODY)
    check_run.font.bold  = True
    check_run.font.color.rgb = RGB_GREEN

    ans_run = p.add_run(f"{letter} — {sanitize(text) if text else ''}")
    ans_run.font.name  = DEFAULT_FONT
    ans_run.font.size  = Pt(FONT_BODY)
    ans_run.font.color.rgb = RGB_GREEN
    return p

def add_explanation(doc, text):
    if not text:
        return
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(4)
    label.paragraph_format.space_after  = Pt(2)
    label.paragraph_format.left_indent  = Inches(0.25)
    set_paragraph_format(label, keep_with_next=True)
    label_run = label.add_run("Explanation:")
    label_run.font.name  = DEFAULT_FONT
    label_run.font.size  = Pt(FONT_SMALL)
    label_run.font.bold  = True
    label_run.font.color.rgb = RGB_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    exp_run = p.add_run(sanitize(text))
    exp_run.font.name  = DEFAULT_FONT
    exp_run.font.size  = Pt(FONT_SMALL)
    exp_run.font.color.rgb = RGB_DARK_TEXT

def add_ite_companion(doc, ite_stem, ite_qid, ite_year):
    """Navy-bordered companion block showing the matched ITE question stem."""
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(10)
    label.paragraph_format.space_after  = Pt(2)
    label.paragraph_format.left_indent  = Inches(0.25)
    add_left_border(label, color=NAVY, size=12)
    add_shading(label, "EBF0F7")
    set_paragraph_format(label, keep_with_next=True)
    label_run = label.add_run(f"Matched ITE Question  ·  {ite_qid}  ·  {ite_year}")
    label_run.font.name  = DEFAULT_FONT
    label_run.font.size  = Pt(FONT_TINY)
    label_run.font.bold  = True
    label_run.font.color.rgb = RGB_NAVY

    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    stem_run = p.add_run(sanitize(ite_stem) if ite_stem else "")
    stem_run.font.name   = DEFAULT_FONT
    stem_run.font.size   = Pt(FONT_SMALL)
    stem_run.font.italic = True
    stem_run.font.color.rgb = RGB_DARK_TEXT

# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    rows = conn.execute(QUERY, (DIST_THRESHOLD,)).fetchall()
    conn.close()
    print(f"Loaded {len(rows)} near-duplicate AAFP questions (dist < {DIST_THRESHOLD})")

    near_dup = sum(1 for r in rows if r["ite_nearest_dist"] < 0.25)
    strong   = sum(1 for r in rows if 0.25 <= r["ite_nearest_dist"] < 0.27)

    doc = new_document()

    add_title(doc, "AAFP Board Review Questions")
    add_subtitle(doc, "ITE Near-Duplicate Set — Semantic Match to ABFM ITE Questions")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    stat_run = p.add_run(
        f"{len(rows)} Questions   |   "
        f"{near_dup} Near-Duplicate (dist < 0.25)   |   "
        f"{strong} Strong Match (0.25–0.27)   |   "
        f"Each question paired with its matched ITE stem"
    )
    stat_run.font.name  = DEFAULT_FONT
    stat_run.font.size  = Pt(FONT_SMALL)
    stat_run.font.color.rgb = RGB_GRAY

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after  = Pt(16)
    note_run = note.add_run("Lower distance = more similar to the ITE question. Sorted closest match first.")
    note_run.font.name   = DEFAULT_FONT
    note_run.font.size   = Pt(FONT_TINY)
    note_run.font.italic = True
    note_run.font.color.rgb = RGB_GRAY

    add_divider(doc)

    for i, row in enumerate(rows, start=1):
        choices = json.loads(row["choices"]) if row["choices"] else []
        correct = (row["correct_letter"] or "").strip().upper()

        add_badge_line(doc, row["ite_nearest_dist"], row["quiz_title"] or "",
                       row["question_number"] or "", row["body_system"] or "",
                       row["ite_nearest_qid"] or "", row["ite_exam_year"] or "")
        add_stem(doc, i, row["stem"] or "")

        for ch in choices:
            add_choice(doc, ch.get("letter", "").strip().upper(), ch.get("text", ""))

        add_correct_answer_line(doc, correct, row["correct_text"] or "")
        add_explanation(doc, row["explanation"])
        add_ite_companion(doc, row["ite_stem"], row["ite_nearest_qid"], row["ite_exam_year"])

        if i < len(rows):
            add_divider(doc)

    for section in doc.sections:
        add_page_number_footer(section)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Saved → {OUT_PATH}")
    print(f"  {len(rows)} questions | {near_dup} near-duplicate | {strong} strong match")

if __name__ == "__main__":
    main()
